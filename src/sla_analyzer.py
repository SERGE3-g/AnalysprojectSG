import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
from datetime import datetime
import os

from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Per chi volesse usare la lettura chunk in futuro:
# from pandas import read_csv


class SLATab(ttk.Frame):
    def __init__(self, parent, current_user, user_role):
        super().__init__(parent)
        self.parent = parent
        self.current_user = current_user
        self.user_role = user_role

        self.log_text = None

        # Variabili per i file
        self.files = {
            'SLA1_IN': {'path': None, 'label': None, 'button': None},
            'SLA4_IN': {'path': None, 'label': None, 'button': None},
            'SLA5_IN': {'path': None, 'label': None, 'button': None},
            'SLA6_IN': {'path': None, 'label': None, 'button': None},
            'SLA1_OUT': {'path': None, 'label': None, 'button': None},
            'SLA2_OUT': {'path': None, 'label': None, 'button': None},
            'SLA3_OUT': {'path': None, 'label': None, 'button': None},
        }

        # Thresholds per ogni SLA
        self.thresholds = {
            'SLA2': 1400,
            'SLA3': 600,
            'SLA4': 800,
            'SLA5': 600,
            'SLA6': 800
        }

        # Creazione widget e layout
        self.create_widgets()
        self.setup_layout()

    def create_widgets(self):
        """Crea tutti i widget necessari per l'interfaccia."""
        # Titolo
        self.title = ttk.Label(
            self,
            text="Analisi SLA",
            font=('Helvetica', 16, 'bold')
        )

        # Frame per la selezione dei file
        self.file_frame = ttk.LabelFrame(self, text="Selezione File", padding=10)

        # Crea label e button di selezione file per ogni chiave self.files
        for sla_key in self.files.keys():
            self.files[sla_key]['label'] = ttk.Label(self.file_frame, text=f"{sla_key}:")
            self.files[sla_key]['button'] = ttk.Button(
                self.file_frame,
                text="Seleziona File",
                command=lambda k=sla_key: self.select_file(k)
            )

        # Frame Opzioni (mese, anno)
        self.options_frame = ttk.LabelFrame(self, text="Opzioni", padding=10)
        self.month_label = ttk.Label(self.options_frame, text="Mese:")
        self.month_var = tk.StringVar(value=datetime.now().strftime("%B"))
        self.month_combo = ttk.Combobox(
            self.options_frame,
            textvariable=self.month_var,
            values=[datetime.strptime(str(i), "%m").strftime("%B") for i in range(1, 13)],
            state="readonly",
            width=15
        )

        self.year_label = ttk.Label(self.options_frame, text="Anno:")
        self.year_var = tk.StringVar(value=str(datetime.now().year))
        self.year_entry = ttk.Entry(
            self.options_frame,
            textvariable=self.year_var,
            width=6
        )

        # Frame bottoni di azione
        self.buttons_frame = ttk.Frame(self)
        self.generate_button = ttk.Button(
            self.buttons_frame,
            text="Genera Report",
            command=self.generate_report
        )
        self.export_csv_btn = ttk.Button(
            self.buttons_frame,
            text="Esporta CSV",
            command=self.export_to_csv
        )
        self.preview_btn = ttk.Button(
            self.buttons_frame,
            text="Anteprima",
            command=self.show_preview
        )
        self.plot_btn = ttk.Button(
            self.buttons_frame,
            text="Mostra Grafico",
            command=self.plot_data
        )

        # Notebook per le varie visualizzazioni
        self.viz_frame = ttk.Notebook(self)

        # Tab per il log
        self.log_frame = ttk.Frame(self.viz_frame)
        self.viz_frame.add(self.log_frame, text="Log")

        # Text + scrollbar per il log
        self.log_text = tk.Text(self.log_frame, height=10, wrap=tk.WORD)
        self.log_scroll = ttk.Scrollbar(
            self.log_frame,
            orient="vertical",
            command=self.log_text.yview
        )
        self.log_text.configure(yscrollcommand=self.log_scroll.set)

        # Tag colore messaggi
        self.log_text.tag_configure("error", foreground="red")
        self.log_text.tag_configure("success", foreground="green")
        self.log_text.tag_configure("info", foreground="blue")
        self.log_text.tag_configure("warning", foreground="orange")

        # Tab per anteprima
        self.preview_frame = ttk.Frame(self.viz_frame)
        self.viz_frame.add(self.preview_frame, text="Anteprima")

        # Treeview + scrollbar per l'anteprima
        self.preview_tree = ttk.Treeview(self.preview_frame)
        self.preview_scroll_y = ttk.Scrollbar(
            self.preview_frame,
            orient="vertical",
            command=self.preview_tree.yview
        )
        self.preview_scroll_x = ttk.Scrollbar(
            self.preview_frame,
            orient="horizontal",
            command=self.preview_tree.xview
        )
        self.preview_tree.configure(
            yscrollcommand=self.preview_scroll_y.set,
            xscrollcommand=self.preview_scroll_x.set
        )

        # Tab per grafico
        self.graph_frame = ttk.Frame(self.viz_frame)
        self.viz_frame.add(self.graph_frame, text="Grafico")

    def log(self, message, level="info"):
        """Aggiunge un messaggio colorato al log."""
        timestamp = datetime.now().strftime('%H:%M:%S')
        self.log_text.insert(tk.END, f"{timestamp}: {message}\n", level)
        self.log_text.see(tk.END)

    def setup_layout(self):
        """Organizza il layout dei widget."""
        # Titolo
        self.title.pack(pady=10)

        # Frame selezione file
        self.file_frame.pack(fill='x', padx=10, pady=5)
        for i, (_, widgets) in enumerate(self.files.items()):
            widgets['label'].grid(row=i, column=0, sticky='w', padx=5, pady=2)
            widgets['button'].grid(row=i, column=1, sticky='w', padx=5, pady=2)

        # Frame opzioni
        self.options_frame.pack(fill='x', padx=10, pady=5)
        self.month_label.pack(side='left', padx=5)
        self.month_combo.pack(side='left', padx=5)
        self.year_label.pack(side='left', padx=5)
        self.year_entry.pack(side='left', padx=5)

        # Frame bottoni
        self.buttons_frame.pack(fill='x', padx=10, pady=5)
        self.generate_button.pack(side='left', padx=5)
        self.export_csv_btn.pack(side='left', padx=5)
        self.preview_btn.pack(side='left', padx=5)
        self.plot_btn.pack(side='left', padx=5)

        # Notebook (log, anteprima, grafico)
        self.viz_frame.pack(fill='both', expand=True, padx=10, pady=5)

        # Log
        self.log_text.pack(side='left', fill='both', expand=True)
        self.log_scroll.pack(side='right', fill='y')

        # Anteprima
        self.preview_tree.pack(side='left', fill='both', expand=True)
        self.preview_scroll_y.pack(side='right', fill='y')
        self.preview_scroll_x.pack(side='bottom', fill='x')

    def select_file(self, sla_key):
        """Gestisce la selezione del file per uno specifico SLA."""
        filename = filedialog.askopenfilename(
            title=f"Seleziona file {sla_key}",
            filetypes=[("CSV files", "*.csv"), ("Tutti i file", "*.*")]
        )
        if filename:
            self.files[sla_key]['path'] = filename
            self.log(f"Selezionato {sla_key}: {filename}", "success")

    def show_preview(self):
        """Mostra l'anteprima del primo file CSV disponibile."""
        try:
            # Pulisce la treeview
            for item in self.preview_tree.get_children():
                self.preview_tree.delete(item)

            # Seleziona il tab dell'anteprima
            self.viz_frame.select(self.preview_frame)

            # Cerca il primo file disponibile
            for file_key, file_data in self.files.items():
                if file_data['path']:
                    df = pd.read_csv(file_data['path'], sep=';', nrows=100)

                    # Configura le colonne
                    self.preview_tree['columns'] = list(df.columns)
                    for col in df.columns:
                        self.preview_tree.heading(col, text=col)
                        self.preview_tree.column(col, width=100)

                    # Inserisce i dati
                    for _, row in df.iterrows():
                        self.preview_tree.insert('', 'end', values=list(row))

                    self.log(f"Anteprima caricata per {file_key}", "success")
                    return

            self.log("Nessun file disponibile per l'anteprima", "warning")

        except Exception as e:
            self.log(f"Errore nell'anteprima: {str(e)}", "error")

    def plot_data(self):
        """Visualizza i dati in un grafico (matplotlib)."""
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

            # Pulisce il frame del grafico
            for widget in self.graph_frame.winfo_children():
                widget.destroy()

            # Seleziona il tab del grafico
            self.viz_frame.select(self.graph_frame)

            # Crea la figura
            fig, ax = plt.subplots(figsize=(10, 6))

            legend_items = []
            for sla_name, threshold in self.thresholds.items():
                file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
                if self.files[file_key]['path']:
                    try:
                        df = pd.read_csv(self.files[file_key]['path'], sep=';')
                        df['Time'] = pd.to_datetime(df['Time'], dayfirst=True)
                        df[sla_name] = pd.to_numeric(df[sla_name], errors='coerce')

                        line = ax.plot(df['Time'], df[sla_name], label=f"SLA {sla_name}", marker='o', markersize=2)
                        ax.axhline(y=threshold, color=line[0].get_color(), linestyle='--', alpha=0.5)
                        legend_items.append(f"{sla_name} (Soglia: {threshold})")

                    except Exception as err:
                        self.log(f"Errore nel plotting di {sla_name}: {str(err)}", "error")
                        continue

            ax.set_title('Analisi SLA', pad=20, fontsize=14)
            ax.set_xlabel('Data/Ora', labelpad=10)
            ax.set_ylabel('Valore', labelpad=10)
            ax.legend(legend_items)
            ax.grid(True, linestyle='--', alpha=0.7)
            fig.autofmt_xdate()
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)

            canvas = FigureCanvasTkAgg(fig, master=self.graph_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill='both', expand=True)

            self.log("Grafico generato con successo", "success")

        except Exception as e:
            self.log(f"Errore nella generazione del grafico: {str(e)}", "error")

    def export_to_csv(self):
        """Esporta i risultati in un unico CSV di riepilogo."""
        try:
            output_path = filedialog.asksaveasfilename(
                defaultextension=".csv",
                filetypes=[("CSV files", "*.csv")],
                initialfile=f"SLA_Analysis_{self.month_var.get()}_{self.year_var.get()}.csv"
            )
            if not output_path:
                return

            all_data = []
            for sla_name, threshold in self.thresholds.items():
                file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
                if self.files[file_key]['path']:
                    df = pd.read_csv(self.files[file_key]['path'], sep=';')
                    df['SLA_Name'] = sla_name
                    df['Threshold'] = threshold
                    df['Source_File'] = os.path.basename(self.files[file_key]['path'])
                    all_data.append(df)

            if all_data:
                final_df = pd.concat(all_data, ignore_index=True)
                final_df['Analysis_Date'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                final_df['Month'] = self.month_var.get()
                final_df['Year'] = self.year_var.get()

                final_df.to_csv(output_path, index=False, sep=';')
                self.log(f"Dati esportati in: {output_path}", "success")
                messagebox.showinfo("Successo", f"Dati esportati in: {output_path}")
            else:
                self.log("Nessun dato da esportare", "warning")

        except Exception as e:
            self.log(f"Errore nell'esportazione CSV: {str(e)}", "error")
            messagebox.showerror("Errore", str(e))

    def generate_report(self):
        """Genera il report completo in formato .docx."""
        try:
            missing_files = [key for key, data in self.files.items() if data['path'] is None]
            if missing_files:
                self.log(f"File mancanti: {', '.join(missing_files)}", "error")
                messagebox.showerror("Errore", "Seleziona tutti i file richiesti")
                return

            self.log("Inizio generazione report...", "info")

            # Directory "reports"
            reports_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
            os.makedirs(reports_dir, exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"SLA_Report_{self.month_var.get()}_{self.year_var.get()}"
            output_path = os.path.join(reports_dir, f"{base_filename}_{timestamp}.docx")

            counter = 1
            while os.path.exists(output_path):
                output_path = os.path.join(reports_dir, f"{base_filename}_{timestamp}_{counter}.docx")
                counter += 1

            # Analisi dati
            results = {}
            for sla_name, threshold in self.thresholds.items():
                file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
                file_path = self.files[file_key]['path']
                try:
                    df = pd.read_csv(file_path, sep=';')
                    df['Time'] = pd.to_datetime(df['Time'], dayfirst=True)
                    df.set_index('Time', inplace=True)

                    results[sla_name] = self.analyze_sla_data(df, threshold, sla_name)
                except Exception as e:
                    self.log(f"Errore nell'analisi di {sla_name}: {str(e)}", "error")
                    raise

            # Genera report Word
            self.generate_docx_report(results, output_path)
            self.log(f"Report generato con successo: {output_path}", "success")
            messagebox.showinfo("Successo", f"Report generato: {output_path}")

            # Tenta di aprire il report
            try:
                os.startfile(output_path)
            except Exception as e:
                self.log(f"Impossibile aprire il report automaticamente: {str(e)}", "warning")

        except Exception as e:
            self.log(f"Errore nella generazione del report: {str(e)}", "error")
            messagebox.showerror("Errore", str(e))

    def generate_docx_report(self, all_results, output_path):
        """Genera il report Word (.docx) finale."""
        try:
            doc = Document()

            # Stile del documento
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
                # Determina il file corretto da utilizzare
                file_key = f'{sla_name}_{"OUT" if sla_name in ["SLA2", "SLA3"] else "IN"}'
                if self.files[file_key]['path'] is None:
                    self.log(f"File mancante per {sla_name}", "error")
                    continue

                try:
                    df = pd.read_csv(self.files[file_key]['path'], sep=';', index_col='Time')
                    df.index = pd.to_datetime(df.index, dayfirst=True)
                    results = self.analyze_sla_data(df, self.thresholds[sla_name], sla_name)

                    section_num = int(sla_name.replace('SLA', ''))
                    doc.add_heading(f'3.{section_num} Rilevazioni {sla_name}', level=2)
                    doc.add_paragraph()

                    # Creazione tabella
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'

                    # Intestazioni
                    headers = [
                        'Data',
                        f'Min con Operatività oltre\nsoglia {self.thresholds[sla_name]} ms',
                        'Min con\nOperatività',
                        '%'
                    ]
                    for i, text in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = text
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Dati giornalieri
                    daily_stats = results['daily_stats']
                    sla_col = sla_name
                    for _, row in daily_stats.iterrows():
                        row_cells = table.add_row().cells
                        row_cells[0].text = row['Date'].strftime('%Y-%m-%d')
                        row_cells[1].text = str(int(row[sla_col]))
                        row_cells[2].text = str(int(row['Time']))

                        total_minutes = int(row['Time'])
                        over_threshold = int(row[sla_col])
                        compliance = (
                            (total_minutes - over_threshold) / total_minutes * 100
                        ) if total_minutes > 0 else 0
                        row_cells[3].text = f"{compliance:.2f}"

                        for cell in row_cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Riga RILEVAZIONE
                    row_cells = table.add_row().cells
                    row_cells[0].text = "RILEVAZIONE"
                    row_cells[1].text = str(results['total_over_threshold'])
                    row_cells[2].text = str(results['total_records'])

                    total_compliance = (
                        (results['total_records'] - results['total_over_threshold']) / results['total_records'] * 100
                    ) if results['total_records'] > 0 else 0
                    row_cells[3].text = f"{total_compliance:.2f}"

                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Conclusioni
                    doc.add_paragraph()
                    threshold_val = 95
                    is_compliant = total_compliance >= threshold_val

                    p = doc.add_paragraph()
                    p.add_run("Il ").italic = True
                    p.add_run("Service Level Agreement").italic = True
                    p.add_run(" si ritiene ")
                    if not is_compliant:
                        p.add_run("non ")
                    p.add_run("rispettato nel periodo di riferimento (")
                    bold_perc = p.add_run(f"{total_compliance:.2f}%")
                    bold_perc.bold = True
                    p.add_run(").")

                    doc.add_paragraph(f"Ciò considerato l'obiettivo previsto del {threshold_val}%.")
                    doc.add_paragraph()

                except Exception as e:
                    self.log(f"Errore nell'elaborazione di {sla_name}: {str(e)}", "error")
                    raise

            doc.save(output_path)
            return True

        except Exception as e:
            self.log(f"Errore nella generazione del report Word: {str(e)}", "error")
            raise

    def set_cell_border(self, tc, top=None, bottom=None):
        """Imposta i bordi di una cella della tabella."""
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

        if top:
            top_element = OxmlElement('w:top')
            top_element.set(qn('w:sz'), str(top['sz']))
            top_element.set(qn('w:val'), top['val'])
            top_element.set(qn('w:color'), top['color'].replace('#', ''))
            tcBorders.append(top_element)

        if bottom:
            bottom_element = OxmlElement('w:bottom')
            bottom_element.set(qn('w:sz'), str(bottom['sz']))
            bottom_element.set(qn('w:val'), bottom['val'])
            bottom_element.set(qn('w:color'), bottom['color'].replace('#', ''))
            tcBorders.append(bottom_element)

    def fast_normalize_time_value(self, value):
        """
        Converte una stringa di tempo in millisecondi (float).
        Versione senza log a ogni riga, per maggiore velocità.
        """
        import re
        import numpy as np

        if pd.isna(value):
            return 0.0  # O np.nan, a scelta

        v = str(value).strip().lower()

        # 1) Se è già un numero (es. "1200" o "1200.5")
        try:
            return float(v)
        except ValueError:
            pass

        # 2) millisecondi con "ms" (es. "647 ms")
        if v.endswith('ms'):
            numeric_str = v.replace('ms', '').strip()
            return float(numeric_str)

        # 3) secondi con "s" (es. "10s")
        if re.match(r'^\d+(\.\d+)?\s*s$', v):
            numeric_str = re.sub(r's$', '', v).strip()
            return float(numeric_str) * 1000.0

        # 4) minuti (es. "2 min", "2mins", "2minutes")
        if re.search(r'(?:mins?|minutes?)$', v):
            numeric_str = re.sub(r'(?:mins?|minutes?)$', '', v).strip()
            return float(numeric_str) * 60000.0

        # Se non riconosciuto, restituiamo 0.0 (o np.nan)
        return 0.0

    def analyze_sla_data(self, df, threshold, sla_name):
        """
        Analizza i dati SLA per produrre statistiche giornaliere in modo più veloce.
        """
        try:
            # Copia DF
            df = df.copy()

            # Se Time è nell'indice, lo resettiamo come colonna
            if df.index.name == 'Time':
                df = df.reset_index()

            # Verifica colonne
            required_columns = ['Time', sla_name]
            for col in required_columns:
                if col not in df.columns:
                    raise KeyError(f"Colonna {col} non trovata nel DataFrame per {sla_name}")

            # Conversione colonna Time in datetime
            df['Time'] = pd.to_datetime(df['Time'], dayfirst=True, errors='coerce')
            if df['Time'].isnull().any():
                self.log(f"Attenzione: ci sono date non parse correttamente per {sla_name}.", "warning")

            # Conversione vettoriale della colonna SLA
            df[sla_name] = df[sla_name].apply(self.fast_normalize_time_value)

            # Aggiunge colonna Date
            df['Date'] = df['Time'].dt.date

            # Calcolo statistiche giornaliere
            daily_stats = df.groupby('Date').agg({
                sla_name: lambda x: (x > threshold).sum(),
                'Time': 'count'
            }).reset_index()

            # Totali
            total_records = int(daily_stats['Time'].sum())
            total_over_threshold = int(daily_stats[sla_name].sum())

            # Dettagli di non conformità (vettoriale)
            non_compliant_df = df[df[sla_name] > threshold].copy()
            non_compliant_df['deviation'] = non_compliant_df[sla_name] - threshold
            non_compliant_details = non_compliant_df[['Time', sla_name, 'deviation']] \
                .rename(columns={sla_name: 'value'}) \
                .to_dict(orient='records')

            results = {
                'daily_stats': daily_stats,
                'total_records': total_records,
                'total_over_threshold': total_over_threshold,
                'non_compliant_details': non_compliant_details,
                'threshold': threshold
            }

            # Log finale di riepilogo
            self.log(f"[{sla_name}] Totale record: {total_records}, Oltre soglia: {total_over_threshold}", "info")
            return results

        except Exception as e:
            self.log(f"Errore nell'analisi di {sla_name}: {str(e)}", "error")
            raise

    def format_table_row(self, row_cells, bold=False, center=True):
        """Formatta una riga della tabella con le impostazioni specificate."""
        for cell in row_cells:
            if center:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if bold:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    def add_table_borders(self, table):
        """Aggiunge i bordi alla tabella."""
        tbl = table._tbl
        borders = tbl.xpath(".//w:tcBorders")
        for border in borders:
            for side in ['top', 'bottom', 'left', 'right']:
                side_element = OxmlElement(f'w:{side}')
                side_element.set(qn('w:sz'), '4')
                side_element.set(qn('w:val'), 'single')
                side_element.set(qn('w:color'), '000000')
                border.append(side_element)

    def process_sla_file(self, sla_name, file_path):
        """Elabora un file SLA e restituisce i dati analizzati."""
        try:
            df = pd.read_csv(file_path, sep=';')
            threshold = self.thresholds[sla_name]
            results = self.analyze_sla_data(df, threshold, sla_name)
            return results
        except Exception as e:
            self.log(f"Errore nell'elaborazione del file {sla_name}: {str(e)}", "error")
            raise

    def initialize_docx_styles(self, doc):
        """Inizializza gli stili del documento Word."""
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        heading_style = doc.styles['Heading 2']
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(13)
        heading_style.font.bold = True

        table_style = doc.styles.add_style('CustomTable', WD_STYLE_TYPE.TABLE)
        table_style.base_style = doc.styles['Table Grid']
        return table_style

    def add_summary_section(self, doc, total_compliance):
        """Aggiunge la sezione di riepilogo al documento."""
        doc.add_paragraph()
        summary = doc.add_heading('Riepilogo Generale', level=1)
        summary.style = doc.styles['Heading 1']

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Indicatore'
        header_cells[1].text = 'Valore'

        row_cells = table.add_row().cells
        row_cells[0].text = 'Compliance Media'
        row_cells[1].text = f"{total_compliance:.2f}%"

        self.format_table_row(header_cells, bold=True)
        self.format_table_row(row_cells)

    def check_required_files(self):
        """Verifica che tutti i file necessari siano stati selezionati."""
        missing = []
        for file_key, file_data in self.files.items():
            if file_data['path'] is None:
                missing.append(file_key)
        return missing

    def validate_month_year(self):
        """Valida il mese e l'anno selezionati."""
        try:
            month = self.month_var.get()
            year = int(self.year_var.get())
            if year < 2000 or year > 2100:
                raise ValueError("Anno non valido")
            return True
        except:
            return False

    def search(self, text):
        """Ricerca testuale nei file SLA caricati e nel log."""
        results = []
        search_text = text.lower()

        # Ricerca nel log
        log_content = self.log_text.get(1.0, tk.END).lower()
        if search_text in log_content:
            relevant_lines = [line for line in log_content.split('\n') if search_text in line.lower()]
            if relevant_lines:
                results.append(("Log SLA", "\n".join(relevant_lines)))

        # Ricerca nei file caricati
        for file_key, file_data in self.files.items():
            if file_data['path']:
                try:
                    df = pd.read_csv(file_data['path'], sep=';')
                    for column in df.columns:
                        matches = df[df[column].astype(str).str.contains(search_text, case=False, na=False)]
                        if not matches.empty:
                            results.append((
                                f"File {file_key} - Colonna {column}",
                                f"Trovate {len(matches)} corrispondenze\n" + matches.head().to_string()
                            ))
                except Exception as e:
                    self.log(f"Errore nella ricerca nel file {file_key}: {str(e)}", "error")

        # Ricerca nelle configurazioni SLA
        try:
            for sla_name, threshold in self.thresholds.items():
                if search_text in sla_name.lower():
                    results.append((
                        "Configurazione SLA",
                        f"SLA: {sla_name}\nSoglia: {threshold}"
                    ))
        except Exception as e:
            self.log(f"Errore nella ricerca configurazioni SLA: {str(e)}", "error")

        if not results:
            results.append(("Info", f"Nessun risultato trovato per: {text}"))

        self.log(f"Ricerca eseguita per: {text} - Trovati {len(results)} risultati", "info")

        return results


def main():
    root = tk.Tk()
    root.title("SLA Analysis")
    app = SLATab(root, current_user="Utente", user_role="admin")
    app.pack(fill='both', expand=True)
    root.mainloop()


if __name__ == "__main__":
    main()
