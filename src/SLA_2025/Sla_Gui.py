import tkinter as tk
from tkinter import ttk, filedialog, messagebox

from django.contrib.admin.checks import check_dependencies
from pip._internal.utils.logging import setup_logging

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
except ImportError:
    # Se tkinterdnd2 non è installato, creiamo un wrapper fittizio
    class TkinterDnD:
        def __init__(self):
            pass

        @staticmethod
        def Tk():
            return tk.Tk()
import os
import sys
import tkinter as tk
from tkinter import ttk, messagebox
import logging
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
except ImportError:
    class TkinterDnD:
        def __init__(self):
            pass

        @staticmethod
        def Tk():
            return tk.Tk()

from datetime import datetime
import threading
from pathlib import Path
import json
import configparser
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from tkinter import colorchooser
import shutil
import webbrowser
from tkcalendar import Calendar
import subprocess


class SLAConfig:
    def __init__(self):
        self.config_file = 'sla_config.ini'
        self.load_config()

    def load_config(self):
        self.config = configparser.ConfigParser()
        if os.path.exists(self.config_file):
            self.config.read(self.config_file)
        else:
            self.create_default_config()

    def create_default_config(self):
        self.config['Thresholds'] = {
            'SLA2': '1400',
            'SLA3': '600',
            'SLA4': '800',
            'SLA5': '600',
            'SLA6': '800'
        }
        self.config['Targets'] = {
            'SLA2': '95',
            'SLA3': '95',
            'SLA4': '95',
            'SLA5': '95',
            'SLA6': '95'
        }
        self.config['Colors'] = {
            'graph_line': '#1f77b4',
            'graph_target': '#d62728',
            'background': '#ffffff'
        }
        self.save_config()

    def save_config(self):
        with open(self.config_file, 'w') as f:
            self.config.write(f)


#class SLAGUI:
    '''def __init__(self, root):
        self.root = root
        print("SLAGUI initialized with root:", root)
        self.root.title("SLA Report Generator Pro")
        self.root.geometry("1200x800")

        # Carica configurazione
        self.config = SLAConfig()

        # Crea menu
        self.create_menu()

        # Crea il notebook per le tab
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Tab principale
        self.main_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.main_tab, text="Report Generator")

        # Tab preview
        self.preview_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.preview_tab, text="Preview")

        # Tab batch processing
        self.batch_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.batch_tab, text="Batch Processing")

        # Tab template manager
        self.template_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.template_tab, text="Template Manager")

        # Inizializza le variabili
        self.init_variables()

        # Crea i widgets per ogni tab
        self.create_main_tab()
        self.create_preview_tab()
        self.create_batch_tab()
        self.create_template_tab()

        # Carica l'ultima sessione
        self.load_last_session()'''

class SLAMonitor:
    def __init__(self, root):
        self.root = root
        self.root.title("SLAMonitor GUI")
        # Aggiungi qui i componenti aggiuntivi della GUI
        tk.Label(root, text="Benvenuto in SLAMonitor!").pack()

        # Load configuration
        self.config = SLAConfig()

        # Initialize variables before creating UI elements
        self.init_variables()

        # Create UI elements
        self.create_menu()
        self.create_notebook()
        self.create_main_tab()
        self.create_preview_tab()
        self.create_batch_tab()
        self.create_template_tab()

        # Load last session at the end
        self.load_session_from_file('last_session.json')

    def init_variables(self):
        """Initialize all necessary variables"""
        self.file_paths = {
            'SLA2': tk.StringVar(),
            'SLA3': tk.StringVar(),
            'SLA4': tk.StringVar(),
            'SLA5': tk.StringVar(),
            'SLA6': tk.StringVar()
        }
        self.output_dir = tk.StringVar()
        self.detailed_graphs = tk.BooleanVar(value=True)
        self.generate_pdf = tk.BooleanVar(value=False)
        self.send_email = tk.BooleanVar(value=False)
        self.progress_var = tk.DoubleVar()
        self.batch_files = []
        self.current_template = None

    def create_notebook(self):
        """Create the notebook for tabs"""
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create frames for each tab
        self.main_tab = ttk.Frame(self.notebook)
        self.preview_tab = ttk.Frame(self.notebook)
        self.batch_tab = ttk.Frame(self.notebook)
        self.template_tab = ttk.Frame(self.notebook)

        # Add tabs to notebook
        self.notebook.add(self.main_tab, text="Report Generator")
        self.notebook.add(self.preview_tab, text="Preview")
        self.notebook.add(self.batch_tab, text="Batch Processing")
        self.notebook.add(self.template_tab, text="Template Manager")

    def load_session_from_file(self, filename):
        """Load session data from a file"""
        try:
            if os.path.exists(filename):
                with open(filename, 'r') as f:
                    session_data = json.load(f)

                for name, path in session_data.get('files', {}).items():
                    if name in self.file_paths:
                        self.file_paths[name].set(path)

                if 'output_dir' in session_data:
                    self.output_dir.set(session_data['output_dir'])

        except Exception as e:
            print(f"Error loading session from {filename}: {e}")

    def save_session_to_file(self, filename):
        """Save current session data to a file"""
        session_data = {
            'files': {name: path.get() for name, path in self.file_paths.items()},
            'output_dir': self.output_dir.get()
        }

        try:
            with open(filename, 'w') as f:
                json.dump(session_data, f)
        except Exception as e:
            print(f"Error saving session to {filename}: {e}")

    def create_menu(self):
        """Create the application menu"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New Report", command=self.clear_fields)
        file_menu.add_command(label="Load Session", command=self.load_session)
        file_menu.add_command(label="Save Session", command=self.save_session)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)

    def create_main_tab(self):
        """Create the main tab content with all controls"""
        # Frame principale con scrollbar
        main_frame = ttk.Frame(self.main_tab)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Canvas e scrollbar per lo scrolling
        canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Sezione periodo di riferimento
        period_frame = ttk.LabelFrame(scrollable_frame, text="Periodo di Riferimento", padding=10)
        period_frame.pack(fill=tk.X, padx=5, pady=5)

        # Calendario per selezione data
        self.cal = Calendar(period_frame,
                            selectmode='day',
                            year=datetime.now().year,
                            month=datetime.now().month)
        self.cal.grid(row=0, column=0, columnspan=2, padx=5, pady=5)

        # Frame selezione file con supporto drag and drop
        files_frame = ttk.LabelFrame(scrollable_frame, text="Selezione File", padding=10)
        files_frame.pack(fill=tk.X, padx=5, pady=5)

        # Creazione campi per ogni tipo di SLA
        for i, (sla_name, path_var) in enumerate(self.file_paths.items()):
            frame = ttk.Frame(files_frame)
            frame.pack(fill=tk.X, padx=5, pady=2)

            ttk.Label(frame, text=f"{sla_name}:").pack(side=tk.LEFT, padx=5)

            # Entry con supporto drag & drop
            entry = ttk.Entry(frame, textvariable=path_var)
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

            # Bottoni per ogni file
            ttk.Button(
                frame,
                text="Sfoglia",
                command=lambda sla=sla_name: self.browse_file(sla)
            ).pack(side=tk.LEFT, padx=2)

            ttk.Button(
                frame,
                text="Preview",
                command=lambda sla=sla_name: self.preview_file(sla)
            ).pack(side=tk.LEFT, padx=2)

            # Supporto drag & drop
            try:
                entry.drop_target_register(DND_FILES)
                entry.dnd_bind('<<Drop>>', lambda e, sla=sla_name: self.drop_file(e, sla))
            except Exception:
                pass  # Ignora se tkinterdnd2 non è disponibile

        # Frame cartella output
        output_frame = ttk.LabelFrame(scrollable_frame, text="Cartella Output", padding=10)
        output_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Entry(
            output_frame,
            textvariable=self.output_dir
        ).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        ttk.Button(
            output_frame,
            text="Sfoglia",
            command=self.browse_output_folder
        ).pack(side=tk.LEFT, padx=5)

        # Opzioni avanzate
        options_frame = ttk.LabelFrame(scrollable_frame, text="Opzioni Avanzate", padding=10)
        options_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Checkbutton(
            options_frame,
            text="Includi grafici dettagliati",
            variable=self.detailed_graphs
        ).pack(side=tk.LEFT, padx=5)

        ttk.Checkbutton(
            options_frame,
            text="Genera PDF",
            variable=self.generate_pdf
        ).pack(side=tk.LEFT, padx=5)

        ttk.Checkbutton(
            options_frame,
            text="Invia email",
            variable=self.send_email
        ).pack(side=tk.LEFT, padx=5)

        # Pulsanti azione
        action_frame = ttk.Frame(scrollable_frame)
        action_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(
            action_frame,
            text="Genera Report",
            command=self.generate_report,
            style='Accent.TButton'
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            action_frame,
            text="Pulisci Campi",
            command=self.clear_fields
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            action_frame,
            text="Anteprima",
            command=self.show_preview
        ).pack(side=tk.LEFT, padx=5)

        # Frame progresso
        progress_frame = ttk.LabelFrame(scrollable_frame, text="Stato", padding=10)
        progress_frame.pack(fill=tk.X, padx=5, pady=5)

        self.progress_bar = ttk.Progressbar(
            progress_frame,
            variable=self.progress_var,
            maximum=100
        )
        self.progress_bar.pack(fill=tk.X, padx=5, pady=5)

        self.status_label = ttk.Label(progress_frame, text="")
        self.status_label.pack(pady=5)

        # Pack finale di canvas e scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Bind per la rotella del mouse
        canvas.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1 * (e.delta / 120)), "units"))

    def create_preview_tab(self):
        """Create the preview tab content with graphs and analysis"""
        # Main frame for preview tab
        preview_frame = ttk.Frame(self.preview_tab)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # Control frame for selection and options
        control_frame = ttk.LabelFrame(preview_frame, text="Controlli", padding=5)
        control_frame.pack(fill=tk.X, padx=5, pady=5)

        # SLA Selection
        ttk.Label(control_frame, text="Seleziona SLA:").pack(side=tk.LEFT, padx=5)
        self.sla_combo = ttk.Combobox(
            control_frame,
            values=list(self.file_paths.keys()),
            state='readonly',
            width=20
        )
        self.sla_combo.pack(side=tk.LEFT, padx=5)
        self.sla_combo.bind('<<ComboboxSelected>>', self.update_preview)

        # Date Range
        ttk.Label(control_frame, text="Periodo:").pack(side=tk.LEFT, padx=5)
        self.period_combo = ttk.Combobox(
            control_frame,
            values=["Ultimo mese", "Ultimi 3 mesi", "Ultimo anno"],
            state='readonly',
            width=15
        )
        self.period_combo.pack(side=tk.LEFT, padx=5)
        self.period_combo.set("Ultimo mese")

        # Refresh button
        ttk.Button(
            control_frame,
            text="Aggiorna",
            command=self.update_preview
        ).pack(side=tk.LEFT, padx=5)

        # Export button
        ttk.Button(
            control_frame,
            text="Esporta Grafico",
            command=self.export_preview
        ).pack(side=tk.RIGHT, padx=5)

        # Create frame for the graph
        self.graph_frame = ttk.LabelFrame(preview_frame, text="Grafico", padding=5)
        self.graph_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create frame for statistics
        self.stats_frame = ttk.LabelFrame(preview_frame, text="Statistiche", padding=5)
        self.stats_frame.pack(fill=tk.X, padx=5, pady=5)

        # Initialize empty labels for statistics
        self.stats_labels = {
            'rows': ttk.Label(self.stats_frame, text="Righe totali: -"),
            'average': ttk.Label(self.stats_frame, text="Media: -"),
            'max': ttk.Label(self.stats_frame, text="Massimo: -"),
            'min': ttk.Label(self.stats_frame, text="Minimo: -")
        }
        for label in self.stats_labels.values():
            label.pack(side=tk.LEFT, padx=10)

    def create_batch_tab(self):
        """Create the batch processing tab for handling multiple files"""
        # Main frame for batch tab
        batch_frame = ttk.Frame(self.batch_tab)
        batch_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # Input folder selection
        folder_frame = ttk.LabelFrame(batch_frame, text="Cartella Input", padding=5)
        folder_frame.pack(fill=tk.X, padx=5, pady=5)

        self.batch_input_var = tk.StringVar()
        ttk.Entry(
            folder_frame,
            textvariable=self.batch_input_var,
            width=50
        ).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        ttk.Button(
            folder_frame,
            text="Sfoglia",
            command=self.browse_batch_folder
        ).pack(side=tk.LEFT, padx=5)

        # Output folder selection
        output_frame = ttk.LabelFrame(batch_frame, text="Cartella Output", padding=5)
        output_frame.pack(fill=tk.X, padx=5, pady=5)

        self.batch_output_var = tk.StringVar()
        ttk.Entry(
            output_frame,
            textvariable=self.batch_output_var,
            width=50
        ).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        ttk.Button(
            output_frame,
            text="Sfoglia",
            command=self.browse_output_folder
        ).pack(side=tk.LEFT, padx=5)

        # File list
        list_frame = ttk.LabelFrame(batch_frame, text="File da Processare", padding=5)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create listbox with scrollbar
        list_container = ttk.Frame(list_frame)
        list_container.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.batch_listbox = tk.Listbox(
            list_container,
            selectmode=tk.EXTENDED,
            height=10
        )
        scrollbar = ttk.Scrollbar(list_container, orient="vertical", command=self.batch_listbox.yview)
        self.batch_listbox.configure(yscrollcommand=scrollbar.set)

        self.batch_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Buttons frame
        button_frame = ttk.Frame(batch_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(
            button_frame,
            text="Processa Tutti",
            command=self.process_batch_all
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            button_frame,
            text="Processa Selezionati",
            command=self.process_batch_selected
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            button_frame,
            text="Rimuovi Selezionati",
            command=self.remove_selected_files
        ).pack(side=tk.LEFT, padx=5)

        # Progress frame
        progress_frame = ttk.LabelFrame(batch_frame, text="Progresso", padding=5)
        progress_frame.pack(fill=tk.X, padx=5, pady=5)

        self.batch_progress = ttk.Progressbar(progress_frame, mode='determinate')
        self.batch_progress.pack(fill=tk.X, padx=5, pady=5)

        self.batch_status_label = ttk.Label(progress_frame, text="Pronto")
        self.batch_status_label.pack(pady=5)

    def create_template_tab(self):
        """Create the template manager tab for handling report templates"""
        # Main frame for template tab
        template_frame = ttk.Frame(self.template_tab)
        template_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # Template selection frame
        selection_frame = ttk.LabelFrame(template_frame, text="Template Attivo", padding=5)
        selection_frame.pack(fill=tk.X, padx=5, pady=5)

        # Template path and browse button
        self.template_path_var = tk.StringVar()
        ttk.Entry(
            selection_frame,
            textvariable=self.template_path_var,
            width=50
        ).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        ttk.Button(
            selection_frame,
            text="Sfoglia",
            command=self.browse_template
        ).pack(side=tk.LEFT, padx=5)

        # Template preview frame
        preview_frame = ttk.LabelFrame(template_frame, text="Anteprima Template", padding=5)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create text widget with scrollbar for template preview
        text_container = ttk.Frame(preview_frame)
        text_container.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.template_text = tk.Text(text_container, wrap=tk.WORD, height=15)
        text_scroll = ttk.Scrollbar(text_container, orient="vertical", command=self.template_text.yview)
        self.template_text.configure(yscrollcommand=text_scroll.set)

        self.template_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        text_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        # Template actions frame
        action_frame = ttk.Frame(template_frame)
        action_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(
            action_frame,
            text="Nuovo Template",
            command=self.new_template
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            action_frame,
            text="Salva Template",
            command=self.save_template
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            action_frame,
            text="Anteprima PDF",
            command=self.preview_template
        ).pack(side=tk.LEFT, padx=5)

        # Variables frame
        variables_frame = ttk.LabelFrame(template_frame, text="Variabili Disponibili", padding=5)
        variables_frame.pack(fill=tk.X, padx=5, pady=5)

        # Create two columns for variables
        left_frame = ttk.Frame(variables_frame)
        right_frame = ttk.Frame(variables_frame)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        right_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Add common variables as labels
        variables = [
            ("${DATA}", "Data del report"),
            ("${PERIODO}", "Periodo di riferimento"),
            ("${SLA_VALORE}", "Valore SLA"),
            ("${SLA_TARGET}", "Target SLA"),
            ("${GRAFICO}", "Grafico performance"),
            ("${TABELLA}", "Tabella dati")
        ]

        for i, (var, desc) in enumerate(variables):
            frame = left_frame if i < len(variables) / 2 else right_frame
            ttk.Label(frame, text=f"{var}: {desc}").pack(anchor=tk.W, padx=5, pady=2)

    def update_preview(self, event=None):
        """Update the preview graph and statistics"""
        selected_sla = self.sla_combo.get()
        if not selected_sla or not self.file_paths[selected_sla].get():
            return

        try:
            # Read the CSV file
            df = pd.read_csv(self.file_paths[selected_sla].get(), sep=';')

            # Clear previous graph
            for widget in self.graph_frame.winfo_children():
                widget.destroy()

            # Create new graph
            fig, ax = plt.subplots(figsize=(10, 6))
            df.plot(kind='line', ax=ax)
            ax.set_title(f'SLA Performance - {selected_sla}')
            ax.grid(True)

            # Add the graph to the frame
            canvas = FigureCanvasTkAgg(fig, self.graph_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

            # Update statistics
            self.stats_labels['rows'].config(text=f"Righe totali: {len(df)}")
            self.stats_labels['average'].config(text=f"Media: {df.iloc[:, 1].mean():.2f}")
            self.stats_labels['max'].config(text=f"Massimo: {df.iloc[:, 1].max():.2f}")
            self.stats_labels['min'].config(text=f"Minimo: {df.iloc[:, 1].min():.2f}")

        except Exception as e:
            messagebox.showerror("Errore", f"Errore nell'aggiornamento del preview: {str(e)}")

    def export_preview(self):
        """Export the current preview graph"""
        if not hasattr(self, 'graph_frame') or not self.graph_frame.winfo_children():
            messagebox.showwarning("Attenzione", "Nessun grafico da esportare")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG files", "*.png"), ("All files", "*.*")]
        )

        if file_path:
            try:
                # Get the current figure and save it
                plt.savefig(file_path)
                messagebox.showinfo("Successo", "Grafico esportato con successo")
            except Exception as e:
                messagebox.showerror("Errore", f"Errore nell'esportazione del grafico: {str(e)}")

    def load_session(self):
        """Load a saved session"""
        try:
            file_path = filedialog.askopenfilename(
                title="Carica Sessione",
                initialdir=os.getcwd(),
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )

            if not file_path:
                return

            self.update_progress(10, "Caricamento sessione in corso...")

            with open(file_path, 'r') as f:
                session_data = json.load(f)

            # Carica i percorsi dei file
            for sla_name, path in session_data.get('files', {}).items():
                if sla_name in self.file_paths:
                    if path and os.path.exists(path):
                        self.file_paths[sla_name].set(path)
                    else:
                        messagebox.showwarning(
                            "Attenzione",
                            f"Il file per {sla_name} non esiste più: {path}"
                        )

            # Carica directory di output
            output_dir = session_data.get('output_dir', '')
            if output_dir and os.path.exists(output_dir):
                self.output_dir.set(output_dir)

            # Carica opzioni
            options = session_data.get('options', {})
            self.detailed_graphs.set(options.get('detailed_graphs', True))
            self.generate_pdf.set(options.get('generate_pdf', False))
            self.send_email.set(options.get('send_email', False))

            # Carica data selezionata
            if 'selected_date' in session_data:
                try:
                    date = datetime.strptime(session_data['selected_date'], '%Y-%m-%d')
                    self.cal.selection_set(date)
                except:
                    pass  # Usa la data corrente se c'è un errore

            self.update_progress(100, "Sessione caricata con successo")
            messagebox.showinfo("Successo", "Sessione caricata correttamente")

        except json.JSONDecodeError:
            messagebox.showerror("Errore", "Il file della sessione è corrotto")
            self.update_progress(0, "Errore nel caricamento della sessione")
        except Exception as e:
            messagebox.showerror("Errore", f"Errore nel caricamento della sessione: {str(e)}")
            self.update_progress(0, "Errore nel caricamento della sessione")

    def save_session(self):
        """Save the current session"""
        try:
            file_path = filedialog.asksaveasfilename(
                title="Salva Sessione",
                initialdir=os.getcwd(),
                defaultextension=".json",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )

            if not file_path:
                return

            self.update_progress(10, "Salvataggio sessione in corso...")

            session_data = {
                'files': {name: path.get() for name, path in self.file_paths.items()},
                'output_dir': self.output_dir.get(),
                'options': {
                    'detailed_graphs': self.detailed_graphs.get(),
                    'generate_pdf': self.generate_pdf.get(),
                    'send_email': self.send_email.get()
                },
                'selected_date': self.cal.get_date(),
                'version': '1.0',  # Versione del formato di salvataggio
                'timestamp': datetime.now().isoformat()
            }

            with open(file_path, 'w') as f:
                json.dump(session_data, f, indent=4)

            self.update_progress(100, "Sessione salvata con successo")
            messagebox.showinfo("Successo", "Sessione salvata correttamente")

            # Salva anche come ultima sessione
            self.save_last_session(session_data)

        except Exception as e:
            messagebox.showerror("Errore", f"Errore nel salvataggio della sessione: {str(e)}")
            self.update_progress(0, "Errore nel salvataggio della sessione")

    def save_last_session(self, session_data=None):
        """Save the last session for auto-recovery"""
        try:
            if session_data is None:
                session_data = {
                    'files': {name: path.get() for name, path in self.file_paths.items()},
                    'output_dir': self.output_dir.get(),
                    'options': {
                        'detailed_graphs': self.detailed_graphs.get(),
                        'generate_pdf': self.generate_pdf.get(),
                        'send_email': self.send_email.get()
                    },
                    'selected_date': self.cal.get_date(),
                    'version': '1.0',
                    'timestamp': datetime.now().isoformat()
                }

            with open('last_session.json', 'w') as f:
                json.dump(session_data, f, indent=4)

        except Exception as e:
            print(f"Errore nel salvataggio dell'ultima sessione: {str(e)}")

    def clear_fields(self):
        """Clear all form fields and reset to default state"""
        try:
            # Chiedi conferma
            if messagebox.askyesno("Conferma", "Vuoi davvero cancellare tutti i campi?"):
                # Reset percorsi file
                for path_var in self.file_paths.values():
                    path_var.set('')

                # Reset directory output
                self.output_dir.set('')

                # Reset opzioni a valori default
                self.detailed_graphs.set(True)
                self.generate_pdf.set(False)
                self.send_email.set(False)

                # Reset progress
                self.progress_var.set(0)
                self.status_label.config(text="")

                # Reset calendario alla data corrente
                today = datetime.now()
                self.cal.selection_set(today)

                # Reset preview se presente
                if hasattr(self, 'graph_frame'):
                    for widget in self.graph_frame.winfo_children():
                        widget.destroy()

                messagebox.showinfo("Completato", "Tutti i campi sono stati ripuliti")

        except Exception as e:
            messagebox.showerror("Errore", f"Errore durante la pulizia dei campi: {str(e)}")

    def show_config_dialog(self):
        """Mostra la finestra di configurazione"""
        config_window = tk.Toplevel(self.root)
        config_window.title("Configurazione")
        config_window.geometry("600x400")

        notebook = ttk.Notebook(config_window)
        notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Tab Soglie
        thresholds_frame = ttk.Frame(notebook)
        notebook.add(thresholds_frame, text="Soglie")

        for sla, threshold in self.config.config['Thresholds'].items():
            frame = ttk.Frame(thresholds_frame)
            frame.pack(fill=tk.X, padx=5, pady=2)

            ttk.Label(frame, text=f"{sla}:").pack(side=tk.LEFT, padx=5)
            entry = ttk.Entry(frame)
            entry.insert(0, threshold)
            entry.pack(side=tk.LEFT, padx=5)
            ttk.Label(frame, text="ms").pack(side=tk.LEFT)

        # Tab Colori
        colors_frame = ttk.Frame(notebook)
        notebook.add(colors_frame, text="Colori")

        for name, color in self.config.config['Colors'].items():
            frame = ttk.Frame(colors_frame)
            frame.pack(fill=tk.X, padx=5, pady=2)

            ttk.Label(frame, text=f"{name}:").pack(side=tk.LEFT, padx=5)
            color_button = ttk.Button(
                frame,
                text="Scegli colore",
                command=lambda n=name: self.choose_color(n)
            )
            color_button.pack(side=tk.LEFT, padx=5)

    def choose_color(self, name):
        """Apre il color picker"""
        color = colorchooser.askcolor(
            color=self.config.config['Colors'][name],
            title=f"Scegli colore per {name}"
        )
        if color[1]:
            self.config.config['Colors'][name] = color[1]
            self.config.save_config()

    def validate_data(self):
        """Valida i dati dei file CSV"""
        validation_window = tk.Toplevel(self.root)
        validation_window.title("Validazione Dati")
        validation_window.geometry("800x600")

        text_widget = tk.Text(validation_window, wrap=tk.WORD)
        text_widget.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        for sla_name, path_var in self.file_paths.items():
            if path_var.get():
                text_widget.insert(tk.END, f"\nValidazione {sla_name}:\n")
                try:
                    df = pd.read_csv(path_var.get(), sep=';')
                    text_widget.insert(tk.END, f"- Righe: {len(df)}\n")
                    text_widget.insert(tk.END, f"- Colonne: {list(df.columns)}\n")
                    text_widget.insert(tk.END, "- Validazione OK\n")
                except Exception as e:
                    text_widget.insert(tk.END, f"- ERRORE: {str(e)}\n")

        text_widget.config(state=tk.DISABLED)

    def show_documentation(self):
        """Apre la documentazione"""
        doc_path = os.path.join(os.path.dirname(__file__), 'docs', 'manual.pdf')
        if os.path.exists(doc_path):
            webbrowser.open(doc_path)
        else:
            messagebox.showinfo("Info", "La documentazione non è disponibile offline.")

    def check_updates(self):
        """Verifica la disponibilità di aggiornamenti"""
        # Qui implementare la logica per il controllo degli aggiornamenti
        messagebox.showinfo("Aggiornamenti", "Nessun aggiornamento disponibile.")

    def show_about(self):
        """Mostra informazioni sul programma"""
        about_text = """
        SLA Report Generator Pro
        Versione 2.0

        Un tool professionale per l'analisi e la reportistica SLA.

        © 2024 Tutti i diritti riservati.
        """
        messagebox.showinfo("Info", about_text)

    def load_session_from_file(self, filename):
                """Load session data from a file"""
                try:
                    if os.path.exists(filename):
                        with open(filename, 'r') as f:
                            session_data = json.load(f)

                        for name, path in session_data.get('files', {}).items():
                            if name in self.file_paths:
                                self.file_paths[name].set(path)

                        if 'output_dir' in session_data:
                            self.output_dir.set(session_data['output_dir'])

                except Exception as e:
                    print(f"Error loading session from {filename}: {e}")

    def save_session_to_file(self, filename):
        """Save current session data to a file"""
        session_data = {
            'files': {name: path.get() for name, path in self.file_paths.items()},
            'output_dir': self.output_dir.get()
        }

        try:
            with open(filename, 'w') as f:
                json.dump(session_data, f)
        except Exception as e:
            print(f"Error saving session to {filename}: {e}")


    def browse_batch_folder(self):
        """Open dialog to select batch input folder"""
        folder = filedialog.askdirectory(title="Seleziona cartella input")
        if folder:
            self.batch_input_var.set(folder)
            self.scan_batch_folder(folder)

    def browse_output_folder(self):
        """Open dialog to select output folder"""
        folder = filedialog.askdirectory(title="Seleziona cartella output")
        if folder:
            self.batch_output_var.set(folder)

    def scan_batch_folder(self, folder):
        """Scan the selected folder for CSV files"""
        self.batch_listbox.delete(0, tk.END)
        try:
            for file in os.listdir(folder):
                if file.endswith('.csv'):
                    self.batch_listbox.insert(tk.END, file)
        except Exception as e:
            messagebox.showerror("Errore", f"Errore nella scansione della cartella: {str(e)}")

    def process_batch_all(self):
        """Process all files in the batch list"""
        if self.batch_listbox.size() == 0:
            messagebox.showinfo("Info", "Nessun file da processare")
            return

        if not self.batch_output_var.get():
            messagebox.showerror("Errore", "Seleziona una cartella di output")
            return

        total_files = self.batch_listbox.size()
        for i in range(total_files):
            filename = self.batch_listbox.get(i)
            self.process_single_file(filename, i, total_files)

    def process_batch_selected(self):
        """Process only selected files in the batch list"""
        selected = self.batch_listbox.curselection()
        if not selected:
            messagebox.showinfo("Info", "Nessun file selezionato")
            return

        if not self.batch_output_var.get():
            messagebox.showerror("Errore", "Seleziona una cartella di output")
            return

        total_selected = len(selected)
        for i, index in enumerate(selected):
            filename = self.batch_listbox.get(index)
            self.process_single_file(filename, i, total_selected)

    def process_single_file(self, filename, current, total):
        """Process a single file from the batch"""
        try:
            input_path = os.path.join(self.batch_input_var.get(), filename)
            output_path = os.path.join(self.batch_output_var.get(), f"processed_{filename}")

            # Update progress
            progress = int((current + 1) / total * 100)
            self.batch_progress['value'] = progress
            self.batch_status_label['text'] = f"Processando {filename}... ({current + 1}/{total})"
            self.batch_tab.update()

            # Process the file (example processing)
            df = pd.read_csv(input_path, sep=';')
            # Add your processing logic here
            df.to_csv(output_path, sep=';', index=False)

            self.batch_status_label['text'] = f"Completato: {filename}"
        except Exception as e:
            self.batch_status_label['text'] = f"Errore nel processare {filename}"
            messagebox.showerror("Errore", f"Errore nel processare {filename}: {str(e)}")

    def remove_selected_files(self):
        """Remove selected files from the batch list"""
        selected = self.batch_listbox.curselection()
        if not selected:
            return

        # Remove in reverse order to maintain correct indices
        for index in sorted(selected, reverse=True):
            self.batch_listbox.delete(index)

    def browse_template(self):
        """Open dialog to select a template file"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if file_path:
            self.template_path_var.set(file_path)
            self.load_template_preview(file_path)

    def load_template_preview(self, template_path):
        """Load and display template preview"""
        try:
            # Here you would implement logic to read the template
            # For now, just display the template path
            self.template_text.delete('1.0', tk.END)
            self.template_text.insert(tk.END, f"Template caricato: {template_path}\n\n")
            self.template_text.insert(tk.END, "Anteprima non disponibile per file DOCX.")
        except Exception as e:
            messagebox.showerror("Errore", f"Errore nel caricamento del template: {str(e)}")

    def new_template(self):
        """Create a new empty template"""
        self.template_path_var.set("")
        self.template_text.delete('1.0', tk.END)

    def save_template(self):
        """Save the current template"""
        if not self.template_path_var.get():
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
            )
            if file_path:
                self.template_path_var.set(file_path)
        try:
            # Here you would implement the actual template saving logic
            messagebox.showinfo("Success", "Template salvato con successo")
        except Exception as e:
            messagebox.showerror("Errore", f"Errore nel salvare il template: {str(e)}")

    def preview_template(self):
        """Preview the current template"""
        if not self.template_path_var.get():
            messagebox.showwarning("Attenzione", "Nessun template selezionato")
            return

        try:
            # Here you would implement the actual preview logic
            # For now, just try to open the file with the default application
            if os.name == 'nt':  # Windows
                os.startfile(self.template_path_var.get())
            else:  # macOS and Linux
                subprocess.call(('open', self.template_path_var.get()))
        except Exception as e:
            messagebox.showerror("Errore", f"Errore nell'apertura del template: {str(e)}")



    def update_preview(self, event):
        """Aggiorna la preview quando viene selezionato un SLA"""
        for widget in self.preview_frame.winfo_children():
            widget.destroy()

        selected_sla = event.widget.get() if event else None
        if not selected_sla or not self.file_paths[selected_sla].get():
            return

        try:
            df = pd.read_csv(self.file_paths[selected_sla].get(), sep=';')

            # Crea il grafico
            fig, ax = plt.subplots(figsize=(10, 6))
            df.plot(kind='line', ax=ax)

            canvas = FigureCanvasTkAgg(fig, self.preview_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

            # Aggiorna le statistiche
            stats_text = f"""
            Numero righe: {len(df)}
            Colonne: {', '.join(df.columns)}
            """
            stats_label = ttk.Label(self.stats_frame, text=stats_text)
            stats_label.pack(padx=5, pady=5)

        except Exception as e:
            messagebox.showerror("Errore", f"Errore nella generazione della preview: {str(e)}")

    def drop_file(self, event, sla_name):
        """Gestisce il drag & drop dei file"""
        file_path = event.data[0]
        if file_path.endswith('.csv'):
            self.file_paths[sla_name].set(file_path)
        else:
            messagebox.showerror("Errore", "Puoi trascinare solo file CSV")

    def generate_report(self):
        """Generate the SLA report with complete processing"""
        # Verifica che tutti i campi necessari siano compilati
        missing_files = [sla for sla, path in self.file_paths.items() if not path.get()]
        if missing_files:
            messagebox.showerror(
                "Errore",
                f"Seleziona i file per: {', '.join(missing_files)}"
            )
            return

        # Verifica cartella output
        if not self.output_dir.get():
            messagebox.showerror(
                "Errore",
                "Seleziona una cartella di output per il report"
            )
            return

        try:
            # Ottieni anno e mese dalla data selezionata
            selected_date = self.cal.get_date()
            date_obj = datetime.strptime(selected_date, '%m/%d/%y')
            year = date_obj.year
            month = date_obj.month
            month_name = date_obj.strftime('%B')

            # Inizializza il processo
            self.update_progress(5, "Inizializzazione processo...")

            # Prepara i file di input
            input_files = {sla_name: path.get() for sla_name, path in self.file_paths.items()}

            # Crea il nome del file di output
            output_file = os.path.join(
                self.output_dir.get(),
                f'Report_SLA_{month_name}_{year}.docx'
            )

            # Leggi e valida i dati
            self.update_progress(20, "Validazione dati in corso...")
            data_frames = {}
            for sla_name, file_path in input_files.items():
                try:
                    df = pd.read_csv(file_path, sep=';')
                    # Validazione base dei dati
                    if df.empty:
                        raise ValueError(f"Il file {sla_name} è vuoto")
                    data_frames[sla_name] = df
                except Exception as e:
                    raise ValueError(f"Errore nella lettura del file {sla_name}: {str(e)}")

            # Analisi dei dati
            self.update_progress(40, "Analisi dati in corso...")
            analyzed_data = {}
            for sla_name, df in data_frames.items():
                try:
                    # Calcola le metriche per ogni SLA
                    metrics = {
                        'total_records': len(df),
                        'average': df.iloc[:, 1].mean(),  # Assume la seconda colonna contiene i valori
                        'max': df.iloc[:, 1].max(),
                        'min': df.iloc[:, 1].min(),
                        'threshold': float(self.config.config['Thresholds'][sla_name]),
                        'target': float(self.config.config['Targets'][sla_name])
                    }
                    analyzed_data[sla_name] = metrics
                except Exception as e:
                    raise ValueError(f"Errore nell'analisi dei dati per {sla_name}: {str(e)}")

            # Generazione grafici
            self.update_progress(60, "Generazione grafici...")
            graphs = {}
            for sla_name, df in data_frames.items():
                try:
                    fig, ax = plt.subplots(figsize=(10, 6))
                    df.plot(kind='line', ax=ax)
                    ax.set_title(f'SLA Performance - {sla_name}')
                    ax.grid(True)
                    graphs[sla_name] = fig
                except Exception as e:
                    raise ValueError(f"Errore nella generazione del grafico per {sla_name}: {str(e)}")

            # Creazione report Word
            self.update_progress(80, "Creazione documento Word...")
            try:
                self.create_word_report(
                    output_file,
                    analyzed_data,
                    graphs,
                    {'year': year, 'month': month, 'month_name': month_name}
                )
            except Exception as e:
                raise ValueError(f"Errore nella creazione del report Word: {str(e)}")

            # Conversione in PDF se richiesto
            if self.generate_pdf.get():
                self.update_progress(90, "Conversione in PDF...")
                pdf_file = output_file.replace('.docx', '.pdf')
                try:
                    self.convert_to_pdf(output_file, pdf_file)
                except Exception as e:
                    messagebox.showwarning(
                        "Attenzione",
                        f"Report Word creato ma conversione PDF fallita: {str(e)}"
                    )

            # Invio email se richiesto
            if self.send_email.get():
                self.update_progress(95, "Invio email...")
                try:
                    self.send_report_email(output_file)
                except Exception as e:
                    messagebox.showwarning(
                        "Attenzione",
                        f"Report creato ma invio email fallito: {str(e)}"
                    )

            # Completamento
            self.update_progress(100, f"Report generato con successo: {output_file}")

            # Mostra messaggio di successo
            messagebox.showinfo(
                "Completato",
                f"Report generato con successo!\nFile salvato in:\n{output_file}"
            )

            # Apri la cartella contenente il file
            if os.name == 'nt':  # Windows
                os.startfile(os.path.dirname(output_file))
            else:  # macOS e Linux
                subprocess.call(['open', os.path.dirname(output_file)])

        except Exception as e:
            self.update_progress(0, f"Errore: {str(e)}")
            messagebox.showerror("Errore", str(e))
            return False

        return True

    def create_word_report(self, output_file, analyzed_data, graphs, period_info):
        """Create Word document with report"""
        from docx import Document
        from docx.shared import Inches

        doc = Document()

        # Aggiungi titolo
        doc.add_heading(f'Report SLA - {period_info["month_name"]} {period_info["year"]}', 0)

        # Aggiungi sommario
        doc.add_paragraph('Sommario dei risultati')
        for sla_name, metrics in analyzed_data.items():
            doc.add_heading(f'Analisi {sla_name}', level=1)
            doc.add_paragraph(f'Totale record: {metrics["total_records"]}')
            doc.add_paragraph(f'Media: {metrics["average"]:.2f}')
            doc.add_paragraph(f'Massimo: {metrics["max"]:.2f}')
            doc.add_paragraph(f'Minimo: {metrics["min"]:.2f}')

            # Salva e inserisci il grafico
            if sla_name in graphs:
                graph_path = f'temp_{sla_name}_graph.png'
                graphs[sla_name].savefig(graph_path)
                doc.add_picture(graph_path, width=Inches(6))
                os.remove(graph_path)  # Pulisci il file temporaneo

        # Salva il documento
        doc.save(output_file)

    def convert_to_pdf(self, docx_file, pdf_file):
        """Convert Word document to PDF using appropriate method based on OS"""
        try:
            if os.name == 'nt':  # Windows
                from docx2pdf import convert
                convert(docx_file, pdf_file)
            else:  # macOS e Linux
                # Usa LibreOffice in modalità headless
                import subprocess
                cmd = ['soffice', '--headless', '--convert-to', 'pdf', '--outdir',
                       os.path.dirname(pdf_file), docx_file]
                subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

            if os.path.exists(pdf_file):
                self.update_progress(92, "PDF creato con successo")
                return True
            else:
                raise Exception("Il file PDF non è stato creato")

        except Exception as e:
            self.update_progress(90, f"Errore nella conversione PDF: {str(e)}")
            raise Exception(f"Errore nella conversione PDF: {str(e)}")

    def send_report_email(self, report_file):
        """Send report via email using configured SMTP settings"""
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        from email.mime.application import MIMEApplication

        try:
            # Carica configurazione email
            email_config = self.load_email_config()

            # Crea il messaggio
            msg = MIMEMultipart()
            msg['From'] = email_config['sender']
            msg['To'] = email_config['recipients']
            msg['Subject'] = f"Report SLA - {datetime.now().strftime('%B %Y')}"

            # Corpo del messaggio
            body = f"""
            Gentili,

            in allegato il report SLA generato il {datetime.now().strftime('%d/%m/%Y')}.

            Cordiali saluti,
            Sistema Report SLA
            """
            msg.attach(MIMEText(body, 'plain'))

            # Allegato
            with open(report_file, 'rb') as f:
                part = MIMEApplication(f.read(), Name=os.path.basename(report_file))
                part['Content-Disposition'] = f'attachment; filename="{os.path.basename(report_file)}"'
                msg.attach(part)

            # Invia email
            with smtplib.SMTP(email_config['smtp_server'], email_config['smtp_port']) as server:
                if email_config.get('use_tls', True):
                    server.starttls()

                if email_config.get('username') and email_config.get('password'):
                    server.login(email_config['username'], email_config['password'])

                server.send_message(msg)

            self.update_progress(97, "Email inviata con successo")
            return True

        except Exception as e:
            self.update_progress(95, f"Errore nell'invio email: {str(e)}")
            raise Exception(f"Errore nell'invio email: {str(e)}")

    def load_email_config(self):
        """Load email configuration from config file"""
        config_file = 'email_config.json'
        try:
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    config = json.load(f)

                # Verifica campi obbligatori
                required_fields = ['sender', 'recipients', 'smtp_server', 'smtp_port']
                missing_fields = [field for field in required_fields if field not in config]

                if missing_fields:
                    raise ValueError(f"Campi mancanti nella configurazione email: {', '.join(missing_fields)}")

                return config
            else:
                # Crea configurazione di default
                default_config = {
                    'sender': 'noreply@yourcompany.com',
                    'recipients': 'recipient@yourcompany.com',
                    'smtp_server': 'smtp.yourcompany.com',
                    'smtp_port': 587,
                    'use_tls': True,
                    'username': '',
                    'password': ''
                }

                with open(config_file, 'w') as f:
                    json.dump(default_config, f, indent=4)

                raise ValueError("Configurazione email non trovata. Creato file di default 'email_config.json'")

        except Exception as e:
            raise Exception(f"Errore nel caricamento della configurazione email: {str(e)}")

    def save_email_config(self, config):
        """Save email configuration to file"""
        try:
            with open('email_config.json', 'w') as f:
                json.dump(config, f, indent=4)
            return True
        except Exception as e:
            raise Exception(f"Errore nel salvataggio della configurazione email: {str(e)}")

    def main(self):
        try:
            root = TkinterDnD.Tk()  # Avvia l'app con TkinterDnD
            app = SLAMonitor(root)  # Creazione dell'istanza di SLAMonitor
            root.mainloop()  # Avvio del ciclo principale della GUI
        except Exception as e:
            print(f"Errore nell'avvio dell'applicazione: {e}")
            raise

    if __name__ == "__main__":
        main()
