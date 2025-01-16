import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import logging
from pathlib import Path
import shutil
import matplotlib.pyplot as plt
import seaborn as sns
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


class SLAConfig:
    """Configurazione degli SLA con soglie e obiettivi"""

    def __init__(self):
        self.thresholds = {
            'SLA2': 1400,  # ms
            'SLA3': 600,  # ms
            'SLA4': 800,  # ms
            'SLA5': 600,  # ms
            'SLA6': 800  # ms
        }
        self.targets = {
            'SLA2': 95,
            'SLA3': 95,
            'SLA4': 95,
            'SLA5': 95,
            'SLA6': 95
        }
        self.directions = {
            'SLA2': 'OUTGOING',
            'SLA3': 'OUTGOING',
            'SLA4': 'INCOMING',
            'SLA5': 'INCOMING',
            'SLA6': 'INCOMING'
        }


class SLAAnalyzer:
    def __init__(self, base_dir=None):
        """
        Inizializza l'analizzatore SLA
        :param base_dir: Directory base per i file di input/output
        """
        self.config = SLAConfig()
        self.base_dir = base_dir or os.getcwd()
        self.setup_logging()

    def setup_logging(self):
        """Configura il logging"""
        log_dir = os.path.join(self.base_dir, 'logs')
        os.makedirs(log_dir, exist_ok=True)

        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(os.path.join(log_dir, 'sla_analyzer.log')),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)

    def create_monthly_folders(self, year, month):
        """
        Crea la struttura delle cartelle per il mese specificato
        """
        month_name = datetime(year, month, 1).strftime('%B').capitalize()
        folders = {
            'input': os.path.join(self.base_dir, f'{year}', f'{month_name}', 'input'),
            'output': os.path.join(self.base_dir, f'{year}', f'{month_name}', 'output'),
            'graphs': os.path.join(self.base_dir, f'{year}', f'{month_name}', 'output', 'graphs')
        }

        for folder in folders.values():
            os.makedirs(folder, exist_ok=True)

        return folders

    def read_csv(self, file_path):
        """
        Legge un file CSV con gestione degli errori e validazione
        """
        try:
            df = pd.read_csv(file_path, sep=';')
            df['Time'] = pd.to_datetime(df['Time'])
            return df
        except Exception as e:
            self.logger.error(f"Errore nella lettura del file {file_path}: {str(e)}")
            raise

    def analyze_sla(self, df, sla_name):
        """
        Analizza i dati SLA con statistiche avanzate
        """
        threshold = self.config.thresholds[sla_name]

        # Raggruppa per giorno
        daily_stats = df.groupby(df['Time'].dt.date).agg({
            sla_name: [
                ('Min con Operatività', 'count'),
                ('Min con Operatività oltre soglia', lambda x: sum(x > threshold)),
                ('Max Response Time', 'max'),
                ('Min Response Time', 'min'),
                ('Avg Response Time', 'mean')
            ]
        })

        daily_stats.columns = daily_stats.columns.droplevel(0)
        daily_stats['%'] = ((daily_stats['Min con Operatività'] - daily_stats['Min con Operatività oltre soglia']) /
                            daily_stats['Min con Operatività'] * 100).round(2)

        # Aggiungi statistiche orarie
        hourly_stats = df.groupby(df['Time'].dt.hour)[sla_name].agg([
            ('count', 'count'),
            ('over_threshold', lambda x: sum(x > threshold)),
            ('compliance', lambda x: (sum(x <= threshold) / len(x) * 100))
        ])

        return daily_stats, hourly_stats

    def create_performance_graphs(self, stats, sla_name, output_folder):
        """
        Crea grafici delle performance
        """
        # Grafico giornaliero
        plt.figure(figsize=(12, 6))
        stats['%'].plot(kind='line', marker='o')
        plt.axhline(y=self.config.targets[sla_name], color='r', linestyle='--',
                    label=f'Target ({self.config.targets[sla_name]}%)')
        plt.title(f'{sla_name} - Performance Giornaliera')
        plt.xlabel('Data')
        plt.ylabel('Compliance (%)')
        plt.grid(True)
        plt.legend()
        plt.tight_layout()

        graph_path = os.path.join(output_folder, f'{sla_name}_daily_performance.png')
        plt.savefig(graph_path)
        plt.close()

        return graph_path

    def add_graph_to_document(self, doc, graph_path):
        """
        Aggiunge un grafico al documento Word
        """
        doc.add_picture(graph_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def create_table_style(self, table):
        """
        Applica uno stile professionale alla tabella
        """
        # Stile intestazione
        header_cells = table.rows[0].cells
        for cell in header_cells:
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.bold = True

        # Stile celle
        for row in table.rows[1:]:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def create_monthly_report(self, year, month, input_files, output_file):
        """
        Crea il report mensile completo
        """
        doc = docx.Document()

        # Stile documento
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Intestazione
        month_name = datetime(year, month, 1).strftime('%B')
        doc.add_heading(f'Report SLA - {month_name} {year}', level=1)

        # Sommario esecutivo
        doc.add_heading('Sommario Esecutivo', level=2)
        summary_table = doc.add_table(rows=1, cols=3)
        summary_table.style = 'Table Grid'
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = 'SLA'
        header_cells[1].text = 'Performance'
        header_cells[2].text = 'Status'

        all_stats = {}

        for sla_name, file_path in input_files.items():
            if sla_name in self.config.thresholds:
                try:
                    df = self.read_csv(file_path)
                    daily_stats, hourly_stats = self.analyze_sla(df, sla_name)
                    all_stats[sla_name] = {'daily': daily_stats, 'hourly': hourly_stats}

                    # Aggiungi sezione al documento
                    doc.add_heading(f'3.{sla_name[-1]} Rilevazioni {sla_name}', level=2)

                    # Aggiungi grafico
                    graph_path = self.create_performance_graphs(daily_stats, sla_name,
                                                                os.path.dirname(output_file))
                    self.add_graph_to_document(doc, graph_path)

                    # Crea tabella principale
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'

                    # Intestazioni
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Data'
                    header_cells[1].text = f'Min con Operatività oltre soglia {self.config.thresholds[sla_name]} ms'
                    header_cells[2].text = 'Min con Operatività'
                    header_cells[3].text = 'Max Response Time (ms)'
                    header_cells[4].text = 'Avg Response Time (ms)'
                    header_cells[5].text = '%'

                    # Aggiungi dati
                    for date, row in daily_stats.iterrows():
                        cells = table.add_row().cells
                        cells[0].text = date.strftime('%Y-%m-%d')
                        cells[1].text = str(int(row['Min con Operatività oltre soglia']))
                        cells[2].text = str(int(row['Min con Operatività']))
                        cells[3].text = f"{row['Max Response Time']:.0f}"
                        cells[4].text = f"{row['Avg Response Time']:.0f}"
                        cells[5].text = f"{row['%']:.2f}"

                    # Totali
                    total_row = table.add_row().cells
                    total_row[0].text = 'RILEVAZIONE'
                    total_row[1].text = str(int(daily_stats['Min con Operatività oltre soglia'].sum()))
                    total_row[2].text = str(int(daily_stats['Min con Operatività'].sum()))
                    total_row[3].text = f"{daily_stats['Max Response Time'].max():.0f}"
                    total_row[4].text = f"{daily_stats['Avg Response Time'].mean():.0f}"

                    total_percentage = ((daily_stats['Min con Operatività'].sum() -
                                         daily_stats['Min con Operatività oltre soglia'].sum()) /
                                        daily_stats['Min con Operatività'].sum() * 100)
                    total_row[5].text = f"{total_percentage:.2f}"

                    self.create_table_style(table)

                    # Aggiungi conclusione
                    rispettato = total_percentage >= self.config.targets[sla_name]
                    conclusion = (f"Il *Service Level Agreement* si ritiene "
                                  f"{'pertanto rispettato' if rispettato else 'non rispettato'} "
                                  f"nel periodo di riferimento (**{total_percentage:.2f}%**).\n\n"
                                  f"Ciò considerato l'obiettivo previsto del {self.config.targets[sla_name]}%.")
                    doc.add_paragraph(conclusion)

                    # Aggiorna sommario
                    row_cells = summary_table.add_row().cells
                    row_cells[0].text = sla_name
                    row_cells[1].text = f"{total_percentage:.2f}%"
                    row_cells[2].text = "✓" if rispettato else "✗"

                except Exception as e:
                    self.logger.error(f"Errore nell'elaborazione di {sla_name}: {str(e)}")
                    continue

        # Salva il documento
        doc.save(output_file)
        self.logger.info(f"Report salvato in: {output_file}")

        return all_stats


def parse_xml(xml_string):
    """Helper function per parsing XML"""
    return OxmlElement(xml_string)


def nsdecls(*prefixes):
    """Helper function per namespace declarations"""
    return ' '.join(['xmlns:{}="{}"'.format(prefix, ns)
                     for prefix, ns in [('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')]
                     if prefix in prefixes])


def main():
    # Crea l'analizzatore
    analyzer = SLAAnalyzer()

    # Ottieni l'anno e il mese corrente
    today = datetime.today()
    year = today.year
    month = today.month - 1  # Analizza il mese precedente
    if month == 0:
        month = 12
        year -= 1

    # Crea le cartelle per il mese
    folders = analyzer.create_monthly_folders(year, month)

    # Definisci i file di input
    month_name = datetime(year, month, 1).strftime('%B')
    input_files = {
        'SLA2': os.path.join(folders['input'], f'SLA2 OUTGOING {month_name} {year}.csv'),
        'SLA3': os.path.join(folders['input'], f'SLA3 OUTGOING {month_name} {year}.csv'),
        'SLA4': os.path.join(folders['input'], f'SLA4 INCOMING {month_name} {year}.csv'),
        'SLA5': os.path.join(folders['input'], f'SLA5 INCOMING {month_name} {year}.csv'),
        'SLA6': os.path.join(folders['input'], f'SLA6 INCOMING {month_name} {year}.csv')
    }

    # Nome del file di output
    output_file = os.path.join(folders['output'], f'Report SLA {month_name} {year}.docx')

    try:
        # Genera il report
        analyzer.create_monthly_report(year, month, input_files, output_file)
        print(f"Report generato con successo: {output_file}")
    except Exception as e:
        print(f"Errore durante la generazione del report: {str(e)}")


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='SLA Report Generator')
    parser.add_argument('--year', type=int, help='Anno per il report (default: anno corrente o precedente per gennaio)')
    parser.add_argument('--month', type=int, help='Mese per il report (1-12, default: mese precedente)')
    parser.add_argument('--base-dir', type=str, help='Directory base per input/output')
    args = parser.parse_args()

    if args.base_dir:
        analyzer = SLAAnalyzer(base_dir=args.base_dir)
    else:
        analyzer = SLAAnalyzer()

    # Se non specificati, usa l'anno e il mese corrente/precedente
    if args.year is None or args.month is None:
        today = datetime.today()
        if args.month is None:
            month = today.month - 1
            if month == 0:
                month = 12
                year = today.year - 1
            else:
                year = today.year
        else:
            month = args.month
            year = args.year or today.year
    else:
        year = args.year
        month = args.month

    # Verifica che il mese sia valido
    if not 1 <= month <= 12:
        print(f"Errore: il mese deve essere compreso tra 1 e 12, non {month}")
        exit(1)

    # Crea le cartelle per il mese
    try:
        folders = analyzer.create_monthly_folders(year, month)

        # Definisci i file di input
        month_name = datetime(year, month, 1).strftime('%B')
        input_files = {
            'SLA2': os.path.join(folders['input'], f'SLA2 OUTGOING {month_name} {year}.csv'),
            'SLA3': os.path.join(folders['input'], f'SLA3 OUTGOING {month_name} {year}.csv'),
            'SLA4': os.path.join(folders['input'], f'SLA4 INCOMING {month_name} {year}.csv'),
            'SLA5': os.path.join(folders['input'], f'SLA5 INCOMING {month_name} {year}.csv'),
            'SLA6': os.path.join(folders['input'], f'SLA6 INCOMING {month_name} {year}.csv')
        }

        # Verifica che tutti i file di input esistano
        missing_files = [f for f in input_files.values() if not os.path.exists(f)]
        if missing_files:
            print("Errore: i seguenti file non sono stati trovati:")
            for f in missing_files:
                print(f"- {f}")
            exit(1)

        # Nome del file di output
        output_file = os.path.join(folders['output'], f'Report SLA {month_name} {year}.docx')

        # Genera il report
        analyzer.create_monthly_report(year, month, input_files, output_file)
        print(f"\nReport generato con successo!")
        print(f"File di output: {output_file}")

    except Exception as e:
        print(f"\nErrore durante la generazione del report: {str(e)}")
        logging.error(f"Errore durante la generazione del report: {str(e)}", exc_info=True)
        exit(1)