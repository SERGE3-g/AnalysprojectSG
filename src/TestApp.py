import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import numpy as np
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import json
import seaborn as sns
from pathlib import Path


class SLAAnalyzer:
    def __init__(self, root):
        self.root = root
        self.root.title("SLA Analyzer Pro")
        self.root.geometry("1200x800")

        # Variabili per i file
        self.files = {
            'SLA1_IN': None,
            'SLA4_IN': None,
            'SLA5_IN': None,
            'SLA6_IN': None,
            'SLA1_OUT': None,
            'SLA2_OUT': None,
            'SLA3_OUT': None
        }

        # Configurazione
        self.config = self.load_config()

        # Crea il menu
        self.create_menu()

        # Crea l'interfaccia principale
        self.create_gui()

        # Crea il notebook per le diverse visualizzazioni
        self.create_notebook()

        # Variabili per i grafici
        self.current_plot = None
        self.plot_data = {}

    def create_menu(self):
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # Menu File
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Load Configuration", command=self.load_config_file)
        file_menu.add_command(label="Save Configuration", command=self.save_config)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)

        # Menu Tools
        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Tools", menu=tools_menu)
        tools_menu.add_command(label="Batch Process", command=self.batch_process)
        tools_menu.add_command(label="Export Charts", command=self.export_charts)
        tools_menu.add_command(label="Data Analysis", command=self.show_data_analysis)

        # Menu Help
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="User Guide", command=self.show_user_guide)
        help_menu.add_command(label="About", command=self.show_about)

    def create_notebook(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.grid(row=5, column=0, columnspan=3, padx=10, pady=5, sticky="nsew")

        # Tab per i grafici
        self.plots_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.plots_frame, text="Grafici")

        # Tab per l'analisi dettagliata
        self.analysis_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.analysis_frame, text="Analisi Dettagliata")

        # Tab per le statistiche
        self.stats_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.stats_frame, text="Statistiche")

    def load_config(self):
        config_file = Path("sla_analyzer_config.json")
        if config_file.exists():
            try:
                with open(config_file, 'r') as f:
                    return json.load(f)
            except:
                return self.get_default_config()
        return self.get_default_config()

    def get_default_config(self):
        return {
            'thresholds': {
                'SLA2': 1400,
                'SLA3': 600,
                'SLA4': 800,
                'SLA5': 600,
                'SLA6': 800
            },
            'target_percentage': 95,
            'last_directory': '',
            'chart_style': 'default',
            'report_template': 'standard'
        }

    def save_config(self):
        try:
            with open('sla_analyzer_config.json', 'w') as f:
                json.dump(self.config, f, indent=4)
            self.log("Configurazione salvata con successo")
        except Exception as e:
            self.log(f"Errore nel salvataggio della configurazione: {str(e)}")

    def create_gui(self):
        # Frame principale con scrollbar
        main_container = ttk.Frame(self.root)
        main_container.grid(row=0, column=0, sticky="nsew")

        # Configurazione del grid
        self.root.grid_rowconfigure(0, weight=1)
        self.root.grid_columnconfigure(0, weight=1)

        # Frame principale
        main_frame = ttk.Frame(main_container, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Stile e tema
        style = ttk.Style()
        style.configure("Title.TLabel", font=('Helvetica', 16, 'bold'))

        # Titolo
        title_label = ttk.Label(main_frame, text="SLA Report Generator Pro", style="Title.TLabel")
        title_label.grid(row=0, column=0, columnspan=3, pady=10)

        # Frame per la selezione dei file
        file_frame = ttk.LabelFrame(main_frame, text="Selezione File", padding="10")
        file_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)

        # Bottoni per la selezione dei file con preview
        row = 0
        for file_key in self.files.keys():
            ttk.Label(file_frame, text=f"{file_key}:").grid(row=row, column=0, sticky=tk.W)
            ttk.Button(file_frame, text="Seleziona File",
                       command=lambda k=file_key: self.select_file(k)).grid(row=row, column=1, padx=5)
            preview_btn = ttk.Button(file_frame, text="Preview",
                                     command=lambda k=file_key: self.preview_file(k))
            preview_btn.grid(row=row, column=2, padx=5)
            row += 1

        # Frame per le opzioni avanzate
        options_frame = ttk.LabelFrame(main_frame, text="Opzioni Avanzate", padding="10")
        options_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)

        # Selezione mese e anno
        ttk.Label(options_frame, text="Mese:").grid(row=0, column=0, sticky=tk.W)
        self.month_var = tk.StringVar(value=datetime.now().strftime("%B"))
        month_combo = ttk.Combobox(options_frame, textvariable=self.month_var,
                                   values=[datetime.strptime(str(i), "%m").strftime("%B") for i in range(1, 13)])
        month_combo.grid(row=0, column=1, padx=5)

        ttk.Label(options_frame, text="Anno:").grid(row=0, column=2, sticky=tk.W)
        self.year_var = tk.StringVar(value=str(datetime.now().year))
        year_entry = ttk.Entry(options_frame, textvariable=self.year_var, width=6)
        year_entry.grid(row=0, column=3, padx=5)

        # Opzioni aggiuntive
        ttk.Label(options_frame, text="Template Report:").grid(row=1, column=0, sticky=tk.W)
        self.template_var = tk.StringVar(value="standard")
        template_combo = ttk.Combobox(options_frame, textvariable=self.template_var,
                                      values=["standard", "dettagliato", "executive"])
        template_combo.grid(row=1, column=1, padx=5)

        ttk.Label(options_frame, text="Stile Grafici:").grid(row=1, column=2, sticky=tk.W)
        self.chart_style_var = tk.StringVar(value="default")
        style_combo = ttk.Combobox(options_frame, textvariable=self.chart_style_var,
                                   values=["default", "dark", "colorful", "minimal"])
        style_combo.grid(row=1, column=3, padx=5)

        # Frame per i bottoni di azione
        action_frame = ttk.Frame(main_frame)
        action_frame.grid(row=3, column=0, columnspan=3, pady=10)

        # Bottoni principali
        ttk.Button(action_frame, text="Genera Report",
                   command=self.generate_report).grid(row=0, column=0, padx=5)
        ttk.Button(action_frame, text="Analisi Rapida",
                   command=self.quick_analysis).grid(row=0, column=1, padx=5)
        ttk.Button(action_frame, text="Pulisci Log",
                   command=self.clear_log).grid(row=0, column=2, padx=5)

        # Area di log con ricerca
        log_frame = ttk.LabelFrame(main_frame, text="Log", padding="5")
        log_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)

        # Campo di ricerca per il log
        self.search_var = tk.StringVar()
        self.search_var.trace('w', self.filter_log)
        ttk.Entry(log_frame, textvariable=self.search_var,
                  placeholder="Cerca nel log...").grid(row=0, column=0, sticky=(tk.W, tk.E))

        # Area di log
        self.log_text = tk.Text(log_frame, height=10, width=70)
        self.log_text.grid(row=1, column=0, pady=5)

        # Scrollbar per l'area di log
        scrollbar = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=self.log_text.yview)
        scrollbar.grid(row=1, column=1, sticky=(tk.N, tk.S))
        self.log_text['yscrollcommand'] = scrollbar.set

    def preview_file(self, file_key):
        if self.files[file_key]:
            try:
                df = pd.read_csv(self.files[file_key], sep=';', nrows=5)
                preview_window = tk.Toplevel(self.root)
                preview_window.title(f"Preview: {file_key}")

                text = tk.Text(preview_window, height=10, width=80)
                text.pack(padx=10, pady=10)
                text.insert(tk.END, df.to_string())
                text.config(state='disabled')
            except Exception as e:
                messagebox.showerror("Error", f"Errore nella preview del file: {str(e)}")
        else:
            messagebox.showwarning("Warning", "Seleziona prima un file")

    def quick_analysis(self):
        """Esegue un'analisi rapida dei dati e mostra i risultati principali"""
        if None in self.files.values():
            messagebox.showwarning("Warning", "Seleziona tutti i file prima dell'analisi")
            return

        try:
            # Crea una finestra per l'analisi rapida
            analysis_window = tk.Toplevel(self.root)
            analysis_window.title("Analisi Rapida")
            analysis_window.geometry("800x600")

            # Notebook per diverse visualizzazioni
            notebook = ttk.Notebook(analysis_window)
            notebook.pack(fill='both', expand=True, padx=10, pady=10)

            # Tab per statistiche generali
            stats_frame = ttk.Frame(notebook)
            notebook.add(stats_frame, text="Statistiche Generali")

            # Calcola e mostra le statistiche principali
            for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
                self.calculate_quick_stats(sla_name, stats_frame)

            # Tab per grafici
            plots_frame = ttk.Frame(notebook)
            notebook.add(plots_frame, text="Grafici")

            # Crea alcuni grafici rapidi
            self.create_quick_plots(plots_frame)

        except Exception as e:
            messagebox.showerror("Error", f"Errore nell'analisi rapida: {str(e)}")

    def calculate_quick_stats(self, sla_name, frame):
        """Calcola statistiche rapide per un singolo SLA"""
        try:
            file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
            df = pd.read_csv(self.files[file_key], sep=';')

            # Calcola statistiche base
            stats = {
                'Media': df[sla_name].mean(),
                'Mediana': df[sla_name].median(),
                'Max': df[sla_name].max(),
                'Min': df[sla_name].min(),
                'Std Dev': df[sla_name].std()
            }

            # Mostra statistiche
            label = ttk.Label(frame, text=f"\nStatistiche {sla_name}:", font=('Helvetica', 10, 'bold'))
            label.pack(anchor='w', padx=10)

            for stat_name, value in stats.items():
                stat_label = ttk.Label(frame, text=f"{stat_name}: {value:.2f}")
                stat_label.pack(anchor='w', padx=20)

        except Exception as e:
            self.log(f"Errore nel calcolo delle statistiche per {sla_name}: {str(e)}")


def create_quick_plots(self, frame):
    """Crea grafici rapidi per l'analisi"""
    try:
        fig = plt.figure(figsize=(10, 6))

        # Crea un grafico per ogni SLA
        for i, sla_name in enumerate(['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6'], 1):
            file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
            df = pd.read_csv(self.files[file_key], sep=';')
            df['Time'] = pd.to_datetime(df['Time'])

            plt.subplot(2, 3, i)
            plt.plot(df['Time'], df[sla_name])
            plt.title(f"{sla_name} Trend")
            plt.xticks(rotation=45)
            plt.tight_layout()

        # Embed il grafico nella GUI
        canvas = FigureCanvasTkAgg(fig, frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    except Exception as e:
        self.log(f"Errore nella creazione dei grafici: {str(e)}")


def filter_log(self, *args):
    """Filtra il contenuto del log in base alla ricerca"""
    search_term = self.search_var.get().lower()

    # Ottieni tutto il testo dal log
    current_text = self.log_text.get("1.0", tk.END)
    self.log_text.tag_remove("highlight", "1.0", tk.END)

    if search_term:
        idx = "1.0"
        while True:
            idx = self.log_text.search(search_term, idx, tk.END, nocase=True)
            if not idx:
                break
            end_idx = f"{idx}+{len(search_term)}c"
            self.log_text.tag_add("highlight", idx, end_idx)
            idx = end_idx

        self.log_text.tag_config("highlight", background="yellow")


def clear_log(self):
    """Pulisce l'area di log"""
    self.log_text.delete(1.0, tk.END)
    self.log("Log pulito")


def batch_process(self):
    """Processa più mesi in batch"""
    try:
        # Chiedi la directory contenente i file da processare
        directory = filedialog.askdirectory(title="Seleziona la directory con i file mensili")
        if not directory:
            return

        # Chiedi il range di mesi da processare
        batch_window = tk.Toplevel(self.root)
        batch_window.title("Processo Batch")
        batch_window.geometry("400x300")

        ttk.Label(batch_window, text="Mese iniziale:").pack(pady=5)
        start_month = ttk.Combobox(batch_window,
                                   values=[datetime.strptime(str(i), "%m").strftime("%B") for i in range(1, 13)])
        start_month.pack(pady=5)

        ttk.Label(batch_window, text="Mese finale:").pack(pady=5)
        end_month = ttk.Combobox(batch_window,
                                 values=[datetime.strptime(str(i), "%m").strftime("%B") for i in range(1, 13)])
        end_month.pack(pady=5)

        ttk.Label(batch_window, text="Anno:").pack(pady=5)
        year = ttk.Entry(batch_window)
        year.pack(pady=5)
        year.insert(0, str(datetime.now().year))

        def start_batch():
            try:
                start_idx = datetime.strptime(start_month.get(), "%B").month
                end_idx = datetime.strptime(end_month.get(), "%B").month
                year_val = int(year.get())

                for month_idx in range(start_idx, end_idx + 1):
                    month_name = datetime.strptime(str(month_idx), "%m").strftime("%B")
                    self.process_month(directory, month_name, year_val)

                messagebox.showinfo("Success", "Processo batch completato!")
                batch_window.destroy()

            except Exception as e:
                messagebox.showerror("Error", f"Errore nel processo batch: {str(e)}")

        ttk.Button(batch_window, text="Avvia Processo", command=start_batch).pack(pady=20)

    except Exception as e:
        self.log(f"Errore nell'avvio del processo batch: {str(e)}")


def process_month(self, directory, month, year):
    """Processa i dati per un mese specifico"""
    try:
        # Imposta i file per il mese corrente
        for sla_type in ['IN', 'OUT']:
            for sla_num in [1, 2, 3, 4, 5, 6]:
                if sla_type == 'IN' and sla_num in [2, 3]:
                    continue

                file_pattern = f"SLA{sla_num}_{sla_type}_{month}_{year}.csv"
                matching_files = [f for f in os.listdir(directory) if f.lower() == file_pattern.lower()]

                if matching_files:
                    self.files[f'SLA{sla_num}_{sla_type}'] = os.path.join(directory, matching_files[0])

        # Genera il report per questo mese
        self.month_var.set(month)
        self.year_var.set(str(year))
        self.generate_report()

    except Exception as e:
        self.log(f"Errore nel processing del mese {month} {year}: {str(e)}")


def show_data_analysis(self):
    """Mostra una finestra di analisi dati avanzata"""
    try:
        analysis_window = tk.Toplevel(self.root)
        analysis_window.title("Analisi Dati Avanzata")
        analysis_window.geometry("1000x800")

        # Notebook per diverse visualizzazioni
        notebook = ttk.Notebook(analysis_window)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)

        # Tab per l'analisi delle correlazioni
        corr_frame = ttk.Frame(notebook)
        notebook.add(corr_frame, text="Correlazioni")
        self.create_correlation_analysis(corr_frame)

        # Tab per l'analisi dei trend
        trend_frame = ttk.Frame(notebook)
        notebook.add(trend_frame, text="Analisi Trend")
        self.create_trend_analysis(trend_frame)

        # Tab per le statistiche avanzate
        stats_frame = ttk.Frame(notebook)
        notebook.add(stats_frame, text="Statistiche Avanzate")
        self.create_advanced_stats(stats_frame)

    except Exception as e:
        self.log(f"Errore nell'apertura dell'analisi dati: {str(e)}")


def create_correlation_analysis(self, frame):
    """Crea l'analisi delle correlazioni tra SLA"""
    try:
        # Prepara i dati
        data = {}
        for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
            file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
            df = pd.read_csv(self.files[file_key], sep=';')
            data[sla_name] = df[sla_name].values

        # Crea DataFrame per la correlazione
        corr_df = pd.DataFrame(data)

        # Calcola e visualizza la matrice di correlazione
        plt.figure(figsize=(10, 8))
        sns.heatmap(corr_df.corr(), annot=True, cmap='coolwarm', center=0)
        plt.title("Matrice di Correlazione tra SLA")

        # Embed il grafico
        canvas = FigureCanvasTkAgg(plt.gcf(), frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    except Exception as e:
        self.log(f"Errore nell'analisi delle correlazioni: {str(e)}")


def create_trend_analysis(self, frame):
    """Crea l'analisi dei trend temporali"""
    try:
        fig = plt.figure(figsize=(12, 8))

        for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
            file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
            df = pd.read_csv(self.files[file_key], sep=';')
            df['Time'] = pd.to_datetime(df['Time'])

            # Calcola la media mobile
            df['MA'] = df[sla_name].rolling(window=100).mean()

            plt.plot(df['Time'], df['MA'], label=sla_name)

        plt.title("Analisi Trend (Media Mobile)")
        plt.legend()
        plt.xticks(rotation=45)
        plt.tight_layout()

        # Embed il grafico
        canvas = FigureCanvasTkAgg(fig, frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    except Exception as e:
        self.log(f"Errore nell'analisi dei trend: {str(e)}")


def create_advanced_stats(self, frame):
    """Crea statistiche avanzate per ogni SLA"""
    try:
        # Crea un text widget per le statistiche
        text = tk.Text(frame, height=20, width=80)
        text.pack(padx=10, pady=10)

        for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
            file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
            df = pd.read_csv(self.files[file_key], sep=';')

            # Calcola statistiche avanzate
            stats = {
                'Media': df[sla_name].mean(),
                'Mediana': df[sla_name].median(),
                'Deviazione Standard': df[sla_name].std(),
                'Skewness': df[sla_name].skew(),
                'Kurtosis': df[sla_name].kurtosis(),
                'Percentile 95': df[sla_name].quantile(0.95),
                'Percentile 99': df[sla_name].quantile(0.99),
                'Autocorrelazione': df[sla_name].autocorr()
            }

            # Aggiungi le statistiche al text widget
            text.insert(tk.END, f"\nStatistiche Avanzate per {sla_name}:\n")
            text.insert(tk.END, "=" * 50 + "\n")
            for stat_name, value in stats.items():
                text.insert(tk.END, f"{stat_name}: {value:.4f}\n")

        text.config(state='disabled')

    except Exception as e:
        self.log(f"Errore nel calcolo delle statistiche avanzate: {str(e)}")


def show_user_guide(self):
    """Mostra la guida utente"""
    guide_window = tk.Toplevel(self.root)
    guide_window.title("Guida Utente")
    guide_window.geometry("800x600")

    text = tk.Text(guide_window, wrap=tk.WORD, padx=10, pady=10)
    text.pack(fill=tk.BOTH, expand=True)

    guide_text = """
    Guida Utente - SLA Analyzer Pro

    1. Selezione dei File
       - Usa i pulsanti "Seleziona File" per caricare i file CSV
       - Puoi visualizzare un'anteprima dei file con il pulsante "Preview"

    2. Generazione Report
       - Seleziona mese e anno
       - Scegli il template del report
       - Clicca su "Genera Report"

    3. Analisi Rapida
       - Fornisce statistiche immediate sui dati
       - Visualizza grafici di base

    4. Analisi Avanzata
       - Accedi a correlazioni tra SLA
       - Visualizza trend temporali
       - Analizza statistiche dettagliate

    5. Processo Batch
       - Elabora più mesi contemporaneamente
       - Seleziona range di date

    6. Configurazione
       - Personalizza soglie SLA
       - Imposta preferenze grafici

    Per supporto aggiuntivo, contatta il team di supporto.
    """

    text.insert(tk.END, guide_text)
    text.config(state='disabled')


def show_about(self):
    """Mostra informazioni sul programma"""
    about_text = """
    SLA Analyzer Pro
    Versione 2.0

    Un'applicazione professionale per l'analisi e il monitoraggio degli SLA.

    Caratteristiche principali:
    - Analisi automatizzata dei dati
    - Generazione report personalizzati
    - Visualizzazioni avanzate
    - Processo batch
    - Analisi statistiche dettagliate

    © 2024 Tutti i diritti riservati
    """

    messagebox.showinfo("About", about_text)

    def export_charts(self):
        """Esporta i grafici in formato immagine"""
        try:
            # Chiedi la directory di esportazione
            export_dir = filedialog.askdirectory(title="Seleziona la directory di esportazione")
            if not export_dir:
                return

            # Crea e salva i vari tipi di grafici
            self.export_trend_charts(export_dir)
            self.export_correlation_charts(export_dir)
            self.export_statistical_charts(export_dir)

            messagebox.showinfo("Success", "Grafici esportati con successo!")

        except Exception as e:
            self.log(f"Errore nell'esportazione dei grafici: {str(e)}")

    def export_trend_charts(self, export_dir):
        """Esporta i grafici dei trend"""
        plt.figure(figsize=(15, 10))

        for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
            file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
            df = pd.read_csv(self.files[file_key], sep=';')
            df['Time'] = pd.to_datetime(df['Time'])

            plt.plot(df['Time'], df[sla_name], label=sla_name)

        plt.title("Trend Analysis - All SLAs")
        plt.legend()
        plt.xticks(rotation=45)
        plt.tight_layout()

        plt.savefig(os.path.join(export_dir, 'trend_analysis.png'), dpi=300, bbox_inches='tight')
        plt.close()

    def export_correlation_charts(self, export_dir):
        """Esporta i grafici delle correlazioni"""
        # Prepara i dati
        data = {}
        for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
            file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
            df = pd.read_csv(self.files[file_key], sep=';')
            data[sla_name] = df[sla_name].values

        corr_df = pd.DataFrame(data)

        plt.figure(figsize=(12, 10))
        sns.heatmap(corr_df.corr(), annot=True, cmap='coolwarm', center=0)
        plt.title("SLA Correlation Matrix")
        plt.tight_layout()

        plt.savefig(os.path.join(export_dir, 'correlation_matrix.png'), dpi=300, bbox_inches='tight')
        plt.close()

    def export_statistical_charts(self, export_dir):
        """Esporta i grafici statistici"""
        # Box plot
        plt.figure(figsize=(15, 8))
        data = []
        labels = []

        for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
            file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
            df = pd.read_csv(self.files[file_key], sep=';')
            data.append(df[sla_name].values)
            labels.append(sla_name)

        plt.boxplot(data, labels=labels)
        plt.title("Statistical Distribution of SLAs")
        plt.ylabel("Values")
        plt.tight_layout()

        plt.savefig(os.path.join(export_dir, 'statistical_distribution.png'), dpi=300, bbox_inches='tight')
        plt.close()

    def analyze_sla_data(self, df, threshold, column_name):
        """Analizza i dati SLA e calcola le statistiche"""
        daily_stats = []

        grouped = df.groupby(df.index.date)
        for date, group in grouped:
            total_minutes = len(group)
            over_threshold = len(group[group[column_name] > threshold])
            percentage = ((total_minutes - over_threshold) / total_minutes) * 100 if total_minutes > 0 else 0

            daily_stats.append({
                'Data': date,
                'Min con Operatività oltre soglia': over_threshold,
                'Min con Operatività': total_minutes,
                '%': percentage
            })

        return pd.DataFrame(daily_stats)

    def generate_docx_report(self, results, output_path):
        """Genera il report in formato DOCX"""
        doc = Document()

        # Stile del documento
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Titolo
        doc.add_heading(f'Report SLA - {self.month_var.get()} {self.year_var.get()}', 0)

        for sla_name, data in results.items():
            # Aggiungi sezione per ogni SLA
            doc.add_heading(f'Rilevazioni {sla_name}', level=1)

            # Crea tabella
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Intestazioni
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Data'
            header_cells[1].text = 'Min con Operatività oltre soglia'
            header_cells[2].text = 'Min con Operatività'
            header_cells[3].text = '%'

            # Dati
            for _, row in data.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['Data'])
                row_cells[1].text = str(row['Min con Operatività oltre soglia'])
                row_cells[2].text = str(row['Min con Operatività'])
                row_cells[3].text = f"{row['%']:.2f}"

            # Riga totale
            row_cells = table.add_row().cells
            row_cells[0].text = 'RILEVAZIONE'
            row_cells[1].text = str(data['Min con Operatività oltre soglia'].sum())
            row_cells[2].text = str(data['Min con Operatività'].sum())

            total_percentage = (data['Min con Operatività'].sum() -
                                data['Min con Operatività oltre soglia'].sum()) / data[
                                   'Min con Operatività'].sum() * 100
            row_cells[3].text = f"{total_percentage:.2f}"

            # Aggiungi valutazione SLA
            doc.add_paragraph()
            threshold = self.config['target_percentage']
            if total_percentage >= threshold:
                doc.add_paragraph(
                    f"Il Service Level Agreement si ritiene pertanto rispettato nel periodo di riferimento "
                    f"({total_percentage:.2f}%)."
                )
            else:
                doc.add_paragraph(
                    f"Il Service Level Agreement si ritiene non rispettato nel periodo di riferimento "
                    f"({total_percentage:.2f}%)."
                )
            doc.add_paragraph(f"Ciò considerato l'obiettivo previsto del {threshold}%.")
            doc.add_paragraph()

        # Salva il documento
        doc.save(output_path)

    def generate_report(self):
        """Genera il report completo"""
        try:
            # Verifica che tutti i file siano stati selezionati
            if None in self.files.values():
                messagebox.showerror("Error", "Seleziona tutti i file richiesti")
                return

            self.log("Inizio generazione report...")

            # Definizione delle soglie per ogni SLA
            thresholds = self.config['thresholds']

            results = {}

            # Analisi per ogni SLA
            for sla_name, threshold in thresholds.items():
                file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"

                # Carica e prepara i dati
                df = pd.read_csv(self.files[file_key], sep=';')
                df['Time'] = pd.to_datetime(df['Time'])
                df.set_index('Time', inplace=True)

                # Analizza i dati
                results[sla_name] = self.analyze_sla_data(df, threshold, sla_name)
                self.log(f"Analisi completata per {sla_name}")

            # Genera il report DOCX
            output_path = f"SLA_Report_{self.month_var.get()}_{self.year_var.get()}.docx"
            self.generate_docx_report(results, output_path)

            self.log(f"Report generato con successo: {output_path}")
            messagebox.showinfo("Success", f"Report generato con successo: {output_path}")

        except Exception as e:
            self.log(f"Errore nella generazione del report: {str(e)}")
            messagebox.showerror("Error", f"Errore nella generazione del report: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = SLAAnalyzer(root)
    root.mainloop()