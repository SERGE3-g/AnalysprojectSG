import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import numpy as np
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class SLAAnalyzer:
    def __init__(self, root):
        self.root = root
        self.root.title("SLA Analyzer")
        self.root.geometry("800x600")

        # Variabili per i file
        self.files = {
            'SLA1_IN': None,
            'SLA4_IN': None,
            'SLA5_IN': None,
            'SLA6_IN': None,
            'SLA1_OUT': None,
            'SLA2_OUT': None,
            'SLA3_OUT': None
        }

        self.create_gui()

    def create_gui(self):
        # Frame principale
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Titolo
        title_label = ttk.Label(main_frame, text="SLA Report Generator", font=('Helvetica', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=3, pady=10)

        # Frame per la selezione dei file
        file_frame = ttk.LabelFrame(main_frame, text="File Selection", padding="10")
        file_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)

        # Bottoni per la selezione dei file
        row = 0
        for file_key in self.files.keys():
            ttk.Label(file_frame, text=f"{file_key}:").grid(row=row, column=0, sticky=tk.W)
            ttk.Button(file_frame, text="Select File",
                       command=lambda k=file_key: self.select_file(k)).grid(row=row, column=1, padx=5)
            row += 1

        # Frame per le opzioni
        options_frame = ttk.LabelFrame(main_frame, text="Options", padding="10")
        options_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)

        # Selezione mese e anno
        ttk.Label(options_frame, text="Month:").grid(row=0, column=0, sticky=tk.W)
        self.month_var = tk.StringVar(value=datetime.now().strftime("%B"))
        month_combo = ttk.Combobox(options_frame, textvariable=self.month_var,
                                   values=[datetime.strptime(str(i), "%m").strftime("%B") for i in range(1, 13)])
        month_combo.grid(row=0, column=1, padx=5)

        ttk.Label(options_frame, text="Year:").grid(row=0, column=2, sticky=tk.W)
        self.year_var = tk.StringVar(value=str(datetime.now().year))
        year_entry = ttk.Entry(options_frame, textvariable=self.year_var, width=6)
        year_entry.grid(row=0, column=3, padx=5)

        # Bottone per generare il report
        ttk.Button(main_frame, text="Generate Report",
                   command=self.generate_report).grid(row=3, column=0, columnspan=3, pady=10)

        # Area di log
        self.log_text = tk.Text(main_frame, height=10, width=70)
        self.log_text.grid(row=4, column=0, columnspan=3, pady=5)

        # Scrollbar per l'area di log
        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.log_text.yview)
        scrollbar.grid(row=4, column=3, sticky=(tk.N, tk.S))
        self.log_text['yscrollcommand'] = scrollbar.set

    def select_file(self, file_key):
        filename = filedialog.askopenfilename(
            title=f"Select {file_key} file",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        if filename:
            self.files[file_key] = filename
            self.log(f"Selected {file_key}: {filename}")

    def log(self, message):
        self.log_text.insert(tk.END, f"{datetime.now().strftime('%H:%M:%S')}: {message}\n")
        self.log_text.see(tk.END)

    def analyze_sla_data(self, df, threshold, column_name):
        # Calcola le statistiche giornaliere
        daily_stats = []

        grouped = df.groupby(df.index.date)
        for date, group in grouped:
            total_minutes = len(group)
            over_threshold = len(group[group[column_name] > threshold])
            percentage = ((total_minutes - over_threshold) / total_minutes) * 100

            daily_stats.append({
                'Data': date,
                'Min con Operatività oltre soglia': over_threshold,
                'Min con Operatività': total_minutes,
                '%': percentage
            })

        return pd.DataFrame(daily_stats)

    def generate_docx_report(self, results, output_path):
        doc = Document()

        # Aggiungi titolo
        doc.add_heading(f'Report SLA - {self.month_var.get()} {self.year_var.get()}', 0)

        for sla_name, data in results.items():
            doc.add_heading(f'Rilevazioni {sla_name}', level=1)

            # Crea tabella
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Intestazioni
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Data'
            header_cells[1].text = 'Min con Operatività oltre soglia'
            header_cells[2].text = 'Min con Operatività'
            header_cells[3].text = '%'

            # Dati
            for _, row in data.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['Data'])
                row_cells[1].text = str(row['Min con Operatività oltre soglia'])
                row_cells[2].text = str(row['Min con Operatività'])
                row_cells[3].text = f"{row['%']:.2f}"

            # Aggiungi riga totale
            row_cells = table.add_row().cells
            row_cells[0].text = 'RILEVAZIONE'
            row_cells[1].text = str(data['Min con Operatività oltre soglia'].sum())
            row_cells[2].text = str(data['Min con Operatività'].sum())
            total_percentage = (data['Min con Operatività'].sum() - data['Min con Operatività oltre soglia'].sum()) / \
                               data['Min con Operatività'].sum() * 100
            row_cells[3].text = f"{total_percentage:.2f}"

            # Aggiungi valutazione SLA
            doc.add_paragraph()
            threshold = 95  # Soglia standard del 95%
            if total_percentage >= threshold:
                doc.add_paragraph(
                    f"Il Service Level Agreement si ritiene pertanto rispettato nel periodo di riferimento ({total_percentage:.2f}%).")
            else:
                doc.add_paragraph(
                    f"Il Service Level Agreement si ritiene non rispettato nel periodo di riferimento ({total_percentage:.2f}%).")
            doc.add_paragraph(f"Ciò considerato l'obiettivo previsto del {threshold}%.")
            doc.add_paragraph()

        doc.save(output_path)

    def generate_report(self):
        try:
            # Verifica che tutti i file siano stati selezionati
            if None in self.files.values():
                messagebox.showerror("Error", "Please select all required files")
                return

            self.log("Starting report generation...")

            # Definizione delle soglie per ogni SLA
            thresholds = {
                'SLA2': 1400,
                'SLA3': 600,
                'SLA4': 800,
                'SLA5': 600,
                'SLA6': 800
            }

            results = {}

            # Analisi per ogni SLA
            for sla_name, threshold in thresholds.items():
                file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"

                # Carica e prepara i dati
                df = pd.read_csv(self.files[file_key], sep=';')
                df['Time'] = pd.to_datetime(df['Time'])
                df.set_index('Time', inplace=True)

                # Analizza i dati
                column_name = sla_name
                results[sla_name] = self.analyze_sla_data(df, threshold, column_name)

            # Genera il report DOCX
            output_path = f"SLA_Report_{self.month_var.get()}_{self.year_var.get()}.docx"
            self.generate_docx_report(results, output_path)

            self.log(f"Report generated successfully: {output_path}")
            messagebox.showinfo("Success", f"Report generated successfully: {output_path}")

        except Exception as e:
            self.log(f"Error generating report: {str(e)}")
            messagebox.showerror("Error", f"Error generating report: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = SLAAnalyzer(root)
    root.mainloop()