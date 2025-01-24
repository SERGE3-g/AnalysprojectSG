import shutil
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

import numpy as np
import pandas as pd
from datetime import datetime, timedelta
import os

from docx import Document
from docx.shared import RGBColor, Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Per chi volesse usare la lettura chunk in futuro:
# from pandas import read_csv


class SLATab(ttk.Frame):
    def __init__(self, parent, current_user, user_role):
        super().__init__(parent)
        self.parent = parent
        self.current_user = current_user
        self.user_role = user_role

        # Inizializza il widget di log
        self.log_text = None
        self.setup_logging()

        # Variabili per i file
        self.files = {
            'SLA1_IN': {'path': None, 'label': None, 'button': None, 'description': 'SLA1 Incoming'},
            'SLA1_OUT': {'path': None, 'label': None, 'button': None, 'description': 'SLA1 Outgoing'},
            'SLA2_OUT': {'path': None, 'label': None, 'button': None, 'description': 'SLA2'},
            'SLA3_OUT': {'path': None, 'label': None, 'button': None, 'description': 'SLA3'},
            'SLA4_IN': {'path': None, 'label': None, 'button': None, 'description': 'SLA4'},
            'SLA5_IN': {'path': None, 'label': None, 'button': None, 'description': 'SLA5'},
            'SLA6_IN': {'path': None, 'label': None, 'button': None, 'description': 'SLA6'},
        }

        # Thresholds per ogni SLA
        self.thresholds = {
            'SLA1_IN': 50,  # 50 transazioni al minuto
            'SLA1_OUT': 50,  # 50 transazioni al minuto
            'SLA2': 1400,
            'SLA3': 600,
            'SLA4': 800,
            'SLA5': 600,
            'SLA6': 800
        }

        # Creazione widget e layout
        self.create_widgets()
        self.setup_layout()

    def setup_logging(self):
        """Inizializza il sistema di logging."""
        # Crea il frame per il log
        self.log_frame = ttk.Frame(self)

        # Crea il widget Text per il log
        self.log_text = tk.Text(self.log_frame, height=10, wrap=tk.WORD)
        self.log_text.tag_configure("error", foreground="red")
        self.log_text.tag_configure("warning", foreground="orange")
        self.log_text.tag_configure("success", foreground="green")
        self.log_text.tag_configure("info", foreground="blue")

        # Aggiungi scrollbar
        self.log_scroll = ttk.Scrollbar(self.log_frame, orient="vertical",
                                        command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=self.log_scroll.set)

    def log(self, message, level="info"):
        """
        Aggiunge un messaggio al log con il livello specificato.

        Args:
            message (str): Il messaggio da loggare
            level (str): Il livello del log (error, warning, success, info)
        """
        if self.log_text is None:
            return

        try:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            log_message = f"{timestamp} - {message}\n"

            self.log_text.insert(tk.END, log_message, level)
            self.log_text.see(tk.END)  # Scroll alla fine

            # Aggiorna l'interfaccia
            self.update_idletasks()

        except Exception as e:
            print(f"Errore nel logging: {str(e)}")
            messagebox.showerror("Errore", f"Errore nel logging: {str(e)}")

    def create_widgets(self):
        """Crea tutti i widget necessari per l'interfaccia."""
        # Titolo
        self.title = ttk.Label(
            self,
            text="Analisi SLA",
            font=('Helvetica', 16, 'bold')
        )

        # Frame per la selezione dei file
        self.file_frame = ttk.LabelFrame(self, text="Selezione File", padding=10)

        # Frame per organizzare i file SLA in due colonne
        self.sla_files_frame = ttk.Frame(self.file_frame)

        # Colonna sinistra per file Incoming
        self.incoming_frame = ttk.LabelFrame(self.sla_files_frame, text="Incoming SLA", padding=5)

        # Colonna destra per file Outgoing
        self.outgoing_frame = ttk.LabelFrame(self.sla_files_frame, text="Outgoing SLA", padding=5)

        # Crea widget per ogni file SLA
        for sla_key, file_info in self.files.items():
            target_frame = self.incoming_frame if '_IN' in sla_key else self.outgoing_frame

            # Frame per ogni SLA
            sla_frame = ttk.Frame(target_frame)

            # Label con descrizione
            file_info['label'] = ttk.Label(sla_frame, text=file_info['description'])

            # Button per selezionare il file
            file_info['button'] = ttk.Button(
                sla_frame,
                text="Seleziona File",
                command=lambda k=sla_key: self.select_file(k)
            )

            # Label per mostrare il path del file
            file_info['path_label'] = ttk.Label(
                sla_frame,
                text="Nessun file selezionato",
                foreground="gray"
            )

            # Preview button
            file_info['preview_button'] = ttk.Button(
                sla_frame,
                text="👁",
                width=3,
                command=lambda k=sla_key: self.preview_file(k)
            )

            # Layout dei widget per questo SLA
            file_info['label'].pack(side='left', padx=5)
            file_info['button'].pack(side='left', padx=5)
            file_info['path_label'].pack(side='left', padx=5)
            file_info['preview_button'].pack(side='left', padx=5)

            # Pack il frame dello SLA
            sla_frame.pack(fill='x', pady=2)

        # Frame Opzioni
        self.options_frame = ttk.LabelFrame(self, text="Opzioni Report", padding=10)

        # Frame per periodo
        self.period_frame = ttk.Frame(self.options_frame)
        self.month_label = ttk.Label(self.period_frame, text="Mese:")
        self.month_var = tk.StringVar(value=datetime.now().strftime("%B"))
        self.month_combo = ttk.Combobox(
            self.period_frame,
            textvariable=self.month_var,
            values=[datetime.strptime(str(i), "%m").strftime("%B") for i in range(1, 13)],
            state="readonly",
            width=15
        )

        self.year_label = ttk.Label(self.period_frame, text="Anno:")
        self.year_var = tk.StringVar(value=str(datetime.now().year))
        self.year_entry = ttk.Entry(
            self.period_frame,
            textvariable=self.year_var,
            width=6
        )

        # Frame per opzioni avanzate
        self.advanced_frame = ttk.LabelFrame(self.options_frame, text="Opzioni Avanzate", padding=5)

        # Checkbox per includere grafici
        self.include_graphs_var = tk.BooleanVar(value=True)
        self.include_graphs_check = ttk.Checkbutton(
            self.advanced_frame,
            text="Includi Grafici",
            variable=self.include_graphs_var
        )

        # Checkbox per analisi dettagliata
        self.detailed_analysis_var = tk.BooleanVar(value=False)
        self.detailed_analysis_check = ttk.Checkbutton(
            self.advanced_frame,
            text="Analisi Dettagliata",
            variable=self.detailed_analysis_var
        )

        # Checkbox per backup
        self.backup_enabled = tk.BooleanVar(value=True)
        self.backup_check = ttk.Checkbutton(
            self.advanced_frame,
            text="Abilita Backup",
            variable=self.backup_enabled
        )

        # Frame per retention dei backup
        self.backup_frame = ttk.Frame(self.advanced_frame)
        self.backup_days_label = ttk.Label(self.backup_frame, text="Giorni retention:")
        self.backup_days = tk.IntVar(value=30)
        self.backup_days_entry = ttk.Entry(
            self.backup_frame,
            textvariable=self.backup_days,
            width=5
        )

        # Frame bottoni azione
        self.buttons_frame = ttk.Frame(self)

        # Bottone principale
        self.generate_button = ttk.Button(
            self.buttons_frame,
            text="Genera Report",
            command=self.generate_report,
            style='Accent.TButton'
        )

        # Altri bottoni
        self.export_csv_btn = ttk.Button(
            self.buttons_frame,
            text="Esporta CSV",
            command=self.export_to_csv
        )

        self.preview_btn = ttk.Button(
            self.buttons_frame,
            text="Anteprima",
            command=self.show_preview
        )

        self.plot_btn = ttk.Button(
            self.buttons_frame,
            text="Mostra Grafico",
            command=self.plot_data
        )

        # Notebook per visualizzazioni
        self.viz_frame = ttk.Notebook(self)

        # Tab per il log
        self.log_frame = ttk.Frame(self.viz_frame)
        self.viz_frame.add(self.log_frame, text="Log")

        # Setup logging se non già fatto
        if self.log_text is None:
            self.setup_logging()

        # Pack log widgets
        self.log_text.pack(side='left', fill='both', expand=True)
        self.log_scroll.pack(side='right', fill='y')

        # Tab per anteprima
        self.preview_frame = ttk.Frame(self.viz_frame)
        self.viz_frame.add(self.preview_frame, text="Anteprima")

        # Treeview per l'anteprima
        self.preview_tree = ttk.Treeview(self.preview_frame)
        self.preview_scroll_y = ttk.Scrollbar(
            self.preview_frame,
            orient="vertical",
            command=self.preview_tree.yview
        )
        self.preview_scroll_x = ttk.Scrollbar(
            self.preview_frame,
            orient="horizontal",
            command=self.preview_tree.xview
        )
        self.preview_tree.configure(
            yscrollcommand=self.preview_scroll_y.set,
            xscrollcommand=self.preview_scroll_x.set
        )

        # Tab per grafico
        self.graph_frame = ttk.Frame(self.viz_frame)
        self.viz_frame.add(self.graph_frame, text="Grafico")

        # Configura stile per il bottone principale
        style = ttk.Style()
        style.configure('Accent.TButton',
                        font=('Helvetica', 10, 'bold'),
                        padding=5)

    def setup_layout(self):
        """Organizza il layout di tutti i widget."""
        # Titolo
        self.title.pack(pady=10)

        # Frame selezione file
        self.file_frame.pack(fill='x', padx=10, pady=5)
        self.sla_files_frame.pack(fill='x', expand=True)

        # Colonne Incoming/Outgoing
        self.incoming_frame.pack(side='left', fill='both', expand=True, padx=5)
        self.outgoing_frame.pack(side='right', fill='both', expand=True, padx=5)

        # Frame opzioni
        self.options_frame.pack(fill='x', padx=10, pady=5)

        # Layout periodo
        self.period_frame.pack(fill='x', pady=5)
        self.month_label.pack(side='left', padx=5)
        self.month_combo.pack(side='left', padx=5)
        self.year_label.pack(side='left', padx=5)
        self.year_entry.pack(side='left', padx=5)

        # Layout opzioni avanzate
        self.advanced_frame.pack(fill='x', pady=5)
        self.include_graphs_check.pack(side='left', padx=10)
        self.detailed_analysis_check.pack(side='left', padx=10)
        self.backup_check.pack(side='left', padx=10)

        # Layout backup frame
        self.backup_frame.pack(side='left', padx=10)
        self.backup_days_label.pack(side='left')
        self.backup_days_entry.pack(side='left', padx=5)

        # Frame bottoni
        self.buttons_frame.pack(fill='x', padx=10, pady=5)
        self.generate_button.pack(side='left', padx=5)
        self.export_csv_btn.pack(side='left', padx=5)
        self.preview_btn.pack(side='left', padx=5)
        self.plot_btn.pack(side='left', padx=5)

        # Area visualizzazione
        self.viz_frame.pack(fill='both', expand=True, padx=10, pady=5)

        # Layout preview
        self.preview_tree.pack(side='left', fill='both', expand=True)
        self.preview_scroll_y.pack(side='right', fill='y')
        self.preview_scroll_x.pack(side='bottom', fill='x')

        # Layout log
        self.log_text.pack(side='left', fill='both', expand=True)
        self.log_scroll.pack(side='right', fill='y')

        # Configura lo stile dei frame
        style = ttk.Style()
        style.configure('TLabelframe', padding=5)
        style.configure('TFrame', padding=2)

        # Inizializza il log
        self.log("Applicazione inizializzata", "info")

    def analyze_sla1_data(self, df, sla_key):
        """
        Analizza i dati per SLA1_IN o SLA1_OUT.

        Args:
            df (pandas.DataFrame): DataFrame con i dati SLA
            sla_key (str): 'SLA1_IN' o 'SLA1_OUT'

        Returns:
            dict: Risultati dell'analisi
        """
        try:
            # Mappa nomi colonne per gestire differenze nei file
            column_mapping = {
                'Failed_Minutes': ['Failed_Minutes', 'Minuti_Fuori_SLA'],
                'Total_Minutes': ['Total_Minutes', 'Minuti_Totali']
            }

            # Verifica e rinomina le colonne in base alla mappatura
            for key, possible_names in column_mapping.items():
                for name in possible_names:
                    if name in df.columns:
                        df = df.rename(columns={name: key})
                        break
                else:
                    error_msg = f"La colonna '{key}' non esiste nel file CSV. Colonne disponibili: {list(df.columns)}"
                    self.log(error_msg, "error")
                    raise ValueError(error_msg)

            # Converte le colonne in numeriche con gestione errori
            for col in column_mapping.keys():
                try:
                    df[col] = pd.to_numeric(df[col], errors='coerce')
                    self.log(f"\nConversione numerica completata per {col}", "success")
                except Exception as e:
                    self.log(f"\nErrore nella conversione numerica di {col}: {str(e)}", "error")
                    raise ValueError(f"Errore nella conversione dei dati per {col}")

            # Rimuovi le righe con valori NaN
            original_len = len(df)
            df = df.dropna(subset=column_mapping.keys())
            if len(df) < original_len:
                self.log(f"\nRimosse {original_len - len(df)} righe con valori mancanti", "warning")

            # Assicura che la colonna 'Date' sia di tipo datetime
            df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
            if df['Date'].isnull().any():
                raise ValueError("La colonna 'Date' contiene valori non validi o mancanti.")

            # Raggruppa per data e calcola le statistiche giornaliere
            daily_stats = df.groupby('Date').agg({
                'Failed_Minutes': 'sum',
                'Total_Minutes': 'first'
            }).reset_index()

            # Calcola la compliance giornaliera
            daily_stats['Compliance'] = ((daily_stats['Total_Minutes'] - daily_stats['Failed_Minutes']) /
                                         daily_stats['Total_Minutes'] * 100)

            # Calcolo totali
            total_minutes = daily_stats['Total_Minutes'].sum()
            total_failed = daily_stats['Failed_Minutes'].sum()
            total_compliance = ((total_minutes - total_failed) / total_minutes * 100) if total_minutes > 0 else 0

            # Trova giorni migliori/peggiori
            worst_day = daily_stats.loc[daily_stats['Failed_Minutes'].idxmax()]
            best_day = daily_stats.loc[daily_stats['Failed_Minutes'].idxmin()]

            # Statistiche aggiuntive
            stats = {
                'media_failed': daily_stats['Failed_Minutes'].mean(),
                'mediana_failed': daily_stats['Failed_Minutes'].median(),
                'std_failed': daily_stats['Failed_Minutes'].std(),
                'giorni_sopra_media': len(daily_stats[daily_stats['Failed_Minutes'] >
                                                      daily_stats['Failed_Minutes'].mean()]),
                'giorni_zero_failed': len(daily_stats[daily_stats['Failed_Minutes'] == 0])
            }

            # Analisi trend
            daily_stats['Failed_Change'] = daily_stats['Failed_Minutes'].diff()
            trend_analysis = {
                'direzione': 'Stabile',
                'variazione_media': daily_stats['Failed_Change'].mean(),
                'giorni_miglioramento': len(daily_stats[daily_stats['Failed_Change'] < 0]),
                'giorni_peggioramento': len(daily_stats[daily_stats['Failed_Change'] > 0])
            }

            if abs(trend_analysis['variazione_media']) > 1:
                trend_analysis['direzione'] = "In peggioramento" if trend_analysis[
                                                                        'variazione_media'] > 0 else "In miglioramento"

            # Preparazione risultati finali
            results = {
                'daily_stats': daily_stats,
                'totals': {
                    'total_minutes': int(total_minutes),
                    'total_failed': int(total_failed),
                    'total_compliance': float(total_compliance),
                    'total_over_threshold': int(total_failed)
                },
                'highlights': {
                    'worst_day': {
                        'date': worst_day['Date'].strftime('%Y-%m-%d'),
                        'failed_minutes': int(worst_day['Failed_Minutes']),
                        'compliance': float(worst_day['Compliance'])
                    },
                    'best_day': {
                        'date': best_day['Date'].strftime('%Y-%m-%d'),
                        'failed_minutes': int(best_day['Failed_Minutes']),
                        'compliance': float(best_day['Compliance'])
                    }
                },
                'statistics': stats,
                'trend': trend_analysis,
                'threshold': self.thresholds[sla_key]
            }

            self.log("\nAnalisi completata con successo", "success")
            return results

        except Exception as e:
            error_msg = f"Errore nell'analisi {sla_key}: {str(e)}"
            self.log(error_msg, "error")
            raise ValueError(error_msg)

    def calculate_sla1_trend(self, daily_stats):
        """
        Calcola il trend per i dati SLA1.

        Args:
            daily_stats (pandas.DataFrame): DataFrame con le statistiche giornaliere

        Returns:
            dict: Dizionario con le informazioni sul trend
        """
        try:
            # Calcola variazioni giornaliere nella compliance
            daily_stats['Compliance_Change'] = daily_stats['Compliance'].diff()

            # Determina la direzione del trend
            avg_change = daily_stats['Compliance_Change'].mean()
            if abs(avg_change) < 0.1:
                trend_direction = "Stabile"
            else:
                trend_direction = "In miglioramento" if avg_change > 0 else "In peggioramento"

            # Calcola volatilità
            volatility = daily_stats['Compliance_Change'].std()
            if volatility < 1:
                stability = "Molto stabile"
            elif volatility < 2:
                stability = "Moderatamente stabile"
            else:
                stability = "Instabile"

            # Identifica pattern
            patterns = []

            # Pattern giorni consecutivi
            consecutive_deterioration = (daily_stats['Compliance_Change'] < 0).rolling(window=3).sum()
            if (consecutive_deterioration >= 3).any():
                patterns.append("Rilevate sequenze di peggioramento consecutivo")

            consecutive_improvement = (daily_stats['Compliance_Change'] > 0).rolling(window=3).sum()
            if (consecutive_improvement >= 3).any():
                patterns.append("Rilevate sequenze di miglioramento consecutivo")

            # Analisi settimanale
            daily_stats['DayOfWeek'] = pd.to_datetime(daily_stats['Date']).dt.day_name()
            weekday_stats = daily_stats.groupby('DayOfWeek')['Compliance'].mean()
            worst_day = weekday_stats.idxmin()
            best_day = weekday_stats.idxmax()

            return {
                'trend_direction': trend_direction,
                'avg_change': float(avg_change),
                'stability': stability,
                'volatility': float(volatility),
                'patterns': patterns,
                'weekday_analysis': {
                    'worst_day': worst_day,
                    'best_day': best_day,
                    'weekday_stats': weekday_stats.to_dict()
                }
            }

        except Exception as e:
            self.log(f"Errore nel calcolo del trend SLA1: {str(e)}", "error")
            return {
                'trend_direction': 'Non calcolabile',
                'avg_change': 0.0,
                'stability': 'Non determinabile',
                'volatility': 0.0,
                'patterns': [],
                'weekday_analysis': {}
            }
    def preview_file(self, sla_key):
        """Mostra un'anteprima del file selezionato."""
        if self.files[sla_key]['path']:
            try:
                df = pd.read_csv(self.files[sla_key]['path'], sep=';', nrows=100)

                # Pulisce la treeview esistente
                for item in self.preview_tree.get_children():
                    self.preview_tree.delete(item)

                # Configura le colonne
                self.preview_tree['columns'] = list(df.columns)
                self.preview_tree['show'] = 'headings'  # Nasconde la colonna dell'id

                for col in df.columns:
                    self.preview_tree.heading(col, text=col)
                    self.preview_tree.column(col, width=100)

                # Inserisce i dati
                for idx, row in df.iterrows():
                    self.preview_tree.insert('', 'end', values=list(row))

                # Seleziona il tab dell'anteprima
                self.viz_frame.select(self.preview_frame)

                self.log(f"Anteprima caricata per {self.files[sla_key]['description']}", "success")

            except Exception as e:
                self.log(f"Errore nel caricamento dell'anteprima: {str(e)}", "error")
        else:
            self.log(f"Nessun file selezionato per {self.files[sla_key]['description']}", "warning")

    def show_preview(self):
        """Mostra l'anteprima del primo file CSV disponibile."""
        try:
            # Pulisce la treeview
            for item in self.preview_tree.get_children():
                self.preview_tree.delete(item)

            # Configura le colonne se non sono già configurate
            if not self.preview_tree['columns']:
                self.preview_tree['columns'] = ('file', 'content')
                self.preview_tree.column('file', width=150)
                self.preview_tree.column('content', width=400)
                self.preview_tree.heading('file', text='File')
                self.preview_tree.heading('content', text='Contenuto')

            # Seleziona il tab dell'anteprima
            self.viz_frame.select(self.preview_frame)

            # Cerca i file disponibili
            preview_data = []
            for file_key, file_data in self.files.items():
                if file_data['path']:
                    try:
                        # Leggi le prime righe del file
                        df = pd.read_csv(file_data['path'], sep=';', nrows=100)

                        # Aggiungi informazioni di base
                        preview_data.append({
                            'file': file_key,
                            'path': file_data['path'],
                            'rows': len(df),
                            'columns': list(df.columns),
                            'sample': df.head(5)
                        })

                    except Exception as e:
                        self.log(f"Errore nella lettura del file {file_key}: {str(e)}", "error")
                        continue

            if not preview_data:
                self.log("Nessun file disponibile per l'anteprima", "warning")
                return

            # Mostra i dati nella treeview
            for data in preview_data:
                # Inserisci nome file e informazioni base
                parent = self.preview_tree.insert('', 'end', text=data['file'],
                                                  values=(data['file'],
                                                          f"Righe: {data['rows']}, Colonne: {len(data['columns'])}"))

                # Inserisci nomi colonne
                cols = self.preview_tree.insert(parent, 'end',
                                                values=('Colonne', ', '.join(data['columns'])))

                # Inserisci prime 5 righe
                for idx, row in data['sample'].iterrows():
                    self.preview_tree.insert(parent, 'end',
                                             values=(f"Riga {idx + 1}", ', '.join(map(str, row.values))))

            # Espandi il primo elemento
            first_item = self.preview_tree.get_children()[0]
            self.preview_tree.item(first_item, open=True)

            self.log("Anteprima caricata con successo", "success")

        except Exception as e:
            self.log(f"Errore nell'anteprima: {str(e)}", "error")
            messagebox.showerror("Errore", f"Errore nell'anteprima: {str(e)}")

    def setup_preview_tree(self):
        """Configura la struttura della treeview per l'anteprima."""
        # Configura le colonne
        self.preview_tree['columns'] = ('file', 'content')
        self.preview_tree.column('#0', width=50)  # Colonna espansione
        self.preview_tree.column('file', width=150)
        self.preview_tree.column('content', width=400)

        # Configura le intestazioni
        self.preview_tree.heading('#0', text='')  # Colonna espansione
        self.preview_tree.heading('file', text='File')
        self.preview_tree.heading('content', text='Contenuto')

        # Configura la scrollbar
        self.preview_scroll_y.config(command=self.preview_tree.yview)
        self.preview_tree.configure(yscrollcommand=self.preview_scroll_y.set)

        # Configura lo stile
        style = ttk.Style()
        style.configure('Treeview', rowheight=25)
        style.configure('Treeview.Heading', font=('Calibri', 10, 'bold'))

    def get_file_info(self, file_path):
        """Ottiene informazioni base su un file CSV."""
        try:
            df = pd.read_csv(file_path, sep=';', nrows=1)
            size = os.path.getsize(file_path) / 1024  # KB
            return {
                'columns': list(df.columns),
                'size': f"{size:.1f} KB",
                'separator': ';',
                'encoding': 'utf-8'  # assumiamo utf-8, potremmo aggiungere detection
            }
        except Exception as e:
            self.log(f"Errore nella lettura del file {file_path}: {str(e)}", "error")
            return None

    def validate_csv(self, file_path):
        """Valida la struttura di un file CSV."""
        try:
            df = pd.read_csv(file_path, sep=';', nrows=5)
            issues = []

            # Verifica colonne necessarie
            required_columns = ['Time']  # aggiungi altre colonne necessarie
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                issues.append(f"Colonne mancanti: {', '.join(missing_columns)}")

            # Verifica tipi di dati
            if 'Time' in df.columns:
                try:
                    pd.to_datetime(df['Time'])
                except:
                    issues.append("La colonna Time non contiene date valide")

            return issues if issues else None

        except Exception as e:
            return [f"Errore nella validazione: {str(e)}"]

    def select_file(self, sla_key):
        """Gestisce la selezione del file per uno specifico SLA."""
        filename = filedialog.askopenfilename(
            title=f"Seleziona file {self.files[sla_key]['description']}",
            filetypes=[("CSV files", "*.csv"), ("Tutti i file", "*.*")]
        )
        if filename:
            self.files[sla_key]['path'] = filename
            # Aggiorna la label del path con il nome del file
            base_filename = os.path.basename(filename)
            self.files[sla_key]['path_label'].config(
                text=base_filename if len(base_filename) < 30
                else f"...{base_filename[-27:]}",
                foreground="green"
            )
            self.log(f"Selezionato {self.files[sla_key]['description']}: {filename}", "success")

    def plot_data(self):
        """Visualizza i dati in un grafico avanzato."""
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
            import matplotlib.dates as mdates

            # Pulisce il frame del grafico
            for widget in self.graph_frame.winfo_children():
                widget.destroy()

            # Crea una figura con subplot per ogni tipo di SLA
            fig = plt.figure(figsize=(12, 8))
            plt.style.use('seaborn')

            # Definisce il layout dei subplot
            gs = fig.add_gridspec(3, 1, hspace=0.3)

            # Raggruppa gli SLA per tipo
            sla_groups = {
                'Throughput': ['SLA1_IN', 'SLA1_OUT'],
                'Response Time': ['SLA2_OUT', 'SLA3_OUT', 'SLA4_IN'],
                'Processing Time': ['SLA5_IN', 'SLA6_IN']
            }

            # Colori per ogni SLA
            colors = plt.cm.Set3(np.linspace(0, 1, 8))

            for idx, (group_name, sla_list) in enumerate(sla_groups.items()):
                ax = fig.add_subplot(gs[idx])

                for sla_key in sla_list:
                    if self.files[sla_key]['path']:
                        try:
                            # Carica i dati
                            df = pd.read_csv(self.files[sla_key]['path'], sep=';')
                            df['Time'] = pd.to_datetime(df['Time'], dayfirst=True)

                            if sla_key.startswith('SLA1'):
                                # Per SLA1, plotta la percentuale di successo
                                success_rate = ((df['Total_Minutes'] - df['Failed_Minutes']) /
                                                df['Total_Minutes'] * 100)
                                ax.plot(df['Time'], success_rate,
                                        label=f"{self.files[sla_key]['description']}",
                                        marker='o', markersize=4)
                            else:
                                # Per altri SLA, plotta i valori diretti
                                sla_col = sla_key.split('_')[0]
                                if sla_col in df.columns:
                                    ax.plot(df['Time'], df[sla_col],
                                            label=f"{self.files[sla_key]['description']}",
                                            marker='o', markersize=4)

                            # Aggiungi linea di threshold
                            if sla_key in self.thresholds:
                                ax.axhline(y=self.thresholds[sla_key],
                                           color='r', linestyle='--', alpha=0.5,
                                           label=f'Threshold ({self.thresholds[sla_key]})')

                        except Exception as e:
                            self.log(f"Errore nel plotting di {sla_key}: {str(e)}", "error")
                            continue

                # Configurazione del subplot
                ax.set_title(f'{group_name} Metrics', pad=10, fontsize=12)
                ax.set_xlabel('Data/Ora', labelpad=10)
                ax.set_ylabel('Valore', labelpad=10)
                ax.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
                ax.grid(True, linestyle='--', alpha=0.7)

                # Formattazione dell'asse x
                ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                ax.xaxis.set_major_locator(mdates.DayLocator(interval=5))
                plt.setp(ax.xaxis.get_majorticklabels(), rotation=45)

                # Rimuovi i bordi superiore e destro
                ax.spines['top'].set_visible(False)
                ax.spines['right'].set_visible(False)

            # Aggiusta il layout per evitare sovrapposizioni
            plt.tight_layout()

            # Crea il canvas per il rendering
            canvas = FigureCanvasTkAgg(fig, master=self.graph_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill='both', expand=True)

            # Aggiungi una toolbar di navigazione
            from matplotlib.backends.backend_tkagg import NavigationToolbar2Tk
            toolbar = NavigationToolbar2Tk(canvas, self.graph_frame)
            toolbar.update()
            canvas.get_tk_widget().pack(fill='both', expand=True)

            # Seleziona il tab del grafico
            self.viz_frame.select(self.graph_frame)

            self.log("Grafici generati con successo", "success")

        except Exception as e:
            self.log(f"Errore nella generazione dei grafici: {str(e)}", "error")
            messagebox.showerror("Errore", f"Errore nella generazione dei grafici: {str(e)}")

    def export_to_csv(self):
        """Esporta i risultati in un unico CSV di riepilogo."""
        try:
            output_path = filedialog.asksaveasfilename(
                defaultextension=".csv",
                filetypes=[("CSV files", "*.csv")],
                initialfile=f"SLA_Analysis_{self.month_var.get()}_{self.year_var.get()}.csv"
            )
            if not output_path:
                return

            all_data = []
            for sla_name, threshold in self.thresholds.items():
                file_key = f"{sla_name}_OUT" if sla_name in ['SLA2', 'SLA3'] else f"{sla_name}_IN"
                if self.files[file_key]['path']:
                    df = pd.read_csv(self.files[file_key]['path'], sep=';')
                    df['SLA_Name'] = sla_name
                    df['Threshold'] = threshold
                    df['Source_File'] = os.path.basename(self.files[file_key]['path'])
                    all_data.append(df)

            if all_data:
                final_df = pd.concat(all_data, ignore_index=True)
                final_df['Analysis_Date'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                final_df['Month'] = self.month_var.get()
                final_df['Year'] = self.year_var.get()

                final_df.to_csv(output_path, index=False, sep=';')
                self.log(f"Dati esportati in: {output_path}", "success")
                messagebox.showinfo("Successo", f"Dati esportati in: {output_path}")
            else:
                self.log("Nessun dato da esportare", "warning")

        except Exception as e:
            self.log(f"Errore nell'esportazione CSV: {str(e)}", "error")
            messagebox.showerror("Errore", str(e))

    def generate_report(self):
        """Genera il report completo in formato .docx."""
        try:
            # Verifica file mancanti
            missing_files = [key for key, data in self.files.items() if data['path'] is None]
            if missing_files:
                self.log(f"File mancanti: {', '.join(missing_files)}", "error")
                messagebox.showerror("Errore", "Seleziona tutti i file richiesti")
                return

            self.log("Inizio generazione report...", "info")

            # Verifica se il backup è abilitato
            if self.backup_enabled.get():
                # Esegui il backup prima di generare il report
                self.create_backup()
                # Pulisci i backup vecchi
                self.cleanup_old_backups(days=self.backup_days.get())

            # Crea directory reports se non esiste
            reports_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
            os.makedirs(reports_dir, exist_ok=True)

            # Genera nome file con timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"SLA_Report_{self.month_var.get()}_{self.year_var.get()}"
            output_path = os.path.join(reports_dir, f"{base_filename}_{timestamp}.docx")

            # Evita sovrascritture
            counter = 1
            while os.path.exists(output_path):
                output_path = os.path.join(reports_dir, f"{base_filename}_{timestamp}_{counter}.docx")
                counter += 1

            # Inizializza documento
            doc = Document()
            self.initialize_document_style(doc)

            # Aggiungi intestazione
            self.add_header(doc)

            # Analisi dati per ogni tipo di SLA
            results = {}

            # Analisi SLA1
            sla1_results = self.analyze_all_sla1()
            if sla1_results:
                self.generate_sla1_section(doc, is_incoming=True)
                self.generate_sla1_section(doc, is_incoming=False)

            # Analisi altri SLA
            for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
                file_key = f'{sla_name}_{"OUT" if sla_name in ["SLA2", "SLA3"] else "IN"}'

                try:
                    df = pd.read_csv(self.files[file_key]['path'], sep=';', index_col='Time')
                    df.index = pd.to_datetime(df.index, dayfirst=True)
                    results[sla_name] = self.analyze_sla_data(df, self.thresholds[sla_name], sla_name)

                    section_num = int(sla_name.replace('SLA', ''))
                    self.generate_sla_section(doc, section_num, sla_name, results[sla_name])

                except Exception as e:
                    self.log(f"Errore nell'analisi di {sla_name}: {str(e)}", "error")
                    raise

            # Aggiungi riepilogo
            self.add_summary_section(doc, results)

            # Aggiungi grafici se richiesto
            if self.include_graphs_var.get():
                self.add_graphs_to_report(doc, results)

            # Aggiungi analisi dettagliata se richiesta
            if self.detailed_analysis_var.get():
                self.add_detailed_analysis(doc, results)

            # Salva il documento
            doc.save(output_path)
            self.log(f"Report generato con successo: {output_path}", "success")

            # Mostra messaggio di successo
            messagebox.showinfo("Successo", f"Report generato: {output_path}")

            # Tenta di aprire il report
            try:
                os.startfile(output_path)
            except Exception as e:
                self.log(f"Impossibile aprire il report automaticamente: {str(e)}", "warning")
                messagebox.showinfo("Info", "Il report è stato generato ma non può essere aperto automaticamente")

        except Exception as e:
            self.log(f"Errore nella generazione del report: {str(e)}", "error")
            messagebox.showerror("Errore", str(e))
            raise

    def initialize_document_style(self, doc):
        """
        Inizializza gli stili del documento con formattazione completa.
        Include stili per titoli, corpo del testo, tabelle, elenchi e sezioni speciali.

        Args:
            doc: Documento Word da formattare
        """
        # Stile Normal (base)
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        # Stile Heading 1
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = 'Calibri'
        heading1_style.font.size = Pt(16)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = RGBColor(0, 51, 102)  # Blu scuro
        heading1_style.paragraph_format.space_before = Pt(24)
        heading1_style.paragraph_format.space_after = Pt(12)
        heading1_style.paragraph_format.keep_with_next = True

        # Stile Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Calibri'
        heading2_style.font.size = Pt(14)
        heading2_style.font.bold = True
        heading2_style.font.color.rgb = RGBColor(51, 51, 153)  # Blu medio
        heading2_style.paragraph_format.space_before = Pt(18)
        heading2_style.paragraph_format.space_after = Pt(6)
        heading2_style.paragraph_format.keep_with_next = True

        # Stile Heading 3
        heading3_style = doc.styles['Heading 3']
        heading3_style.font.name = 'Calibri'
        heading3_style.font.size = Pt(12)
        heading3_style.font.bold = True
        heading3_style.font.color.rgb = RGBColor(68, 84, 106)  # Grigio blu
        heading3_style.paragraph_format.space_before = Pt(12)
        heading3_style.paragraph_format.space_after = Pt(6)

        # Stile per tabelle
        if 'Custom Table' not in doc.styles:
            table_style = doc.styles.add_style('Custom Table', WD_STYLE_TYPE.PARAGRAPH)
            table_style.base_style = doc.styles['Normal']
            table_style.font.size = Pt(10)
            table_style.paragraph_format.space_after = Pt(0)
            table_style.paragraph_format.line_spacing = 1

        # Stile per caption tabelle e figure
        if 'Caption Style' not in doc.styles:
            caption_style = doc.styles.add_style('Caption Style', WD_STYLE_TYPE.PARAGRAPH)
            caption_style.base_style = doc.styles['Normal']
            caption_style.font.size = Pt(9)
            caption_style.font.italic = True
            caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_style.paragraph_format.space_before = Pt(6)
            caption_style.paragraph_format.space_after = Pt(12)

        # Configura stili tabella predefiniti
        table_styles = ['Table Grid', 'Light Shading', 'Light List', 'Light Grid']
        for style_name in table_styles:
            if style_name in doc.styles:
                table_style = doc.styles[style_name]
                table_style.font.size = Pt(10)
                table_style.font.name = 'Calibri'

        # Crea sezione per il documento
        section = doc.sections[0]
        section.page_height = Mm(297)  # A4
        section.page_width = Mm(210)  # A4
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        # Aggiungi header e footer
        header = section.header
        header.is_linked_to_previous = False
        footer = section.footer
        footer.is_linked_to_previous = False

        # Aggiungi numeri di pagina nel footer usando il metodo corretto
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metodo corretto per aggiungere la numerazione delle pagine
        run = footer_paragraph.add_run()
        run.text = 'Pagina '

        # Aggiungi campo PAGE
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char1)

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'PAGE'
        run._r.append(instr_text)

        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char2)

        run = footer_paragraph.add_run()
        run.text = ' di '

        # Aggiungi campo NUMPAGES
        run = footer_paragraph.add_run()
        fld_char3 = OxmlElement('w:fldChar')
        fld_char3.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char3)

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'NUMPAGES'
        run._r.append(instr_text)

        fld_char4 = OxmlElement('w:fldChar')
        fld_char4.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char4)

    def add_header(self, doc):
        """Aggiunge l'intestazione al documento."""
        # Titolo principale
        title = doc.add_heading('Report Analisi SLA', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informazioni periodo
        period = doc.add_paragraph()
        period.add_run(f'Periodo: {self.month_var.get()} {self.year_var.get()}')
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data generazione
        gen_date = doc.add_paragraph()
        gen_date.add_run(f'Generato il: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spazio vuoto

    def add_graphs_to_report(self, doc, results):
        """Aggiunge i grafici al report."""
        doc.add_heading('Grafici', level=1)

        # Crea i grafici
        try:
            import matplotlib.pyplot as plt
            import io

            for sla_name, data in results.items():
                plt.figure(figsize=(10, 6))
                # ... codice per creare il grafico specifico per ogni SLA ...

                # Salva il grafico come immagine
                img_stream = io.BytesIO()
                plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
                img_stream.seek(0)

                # Aggiungi l'immagine al documento
                doc.add_picture(img_stream, width=Inches(6))
                doc.add_paragraph(f'Grafico {sla_name}').alignment = WD_ALIGN_PARAGRAPH.CENTER

                plt.close()

        except Exception as e:
            self.log(f"Errore nella generazione dei grafici: {str(e)}", "error")
            doc.add_paragraph("Impossibile generare i grafici").italic = True

    def add_detailed_analysis(self, doc, results):
        """
        Aggiunge l'analisi dettagliata al report con statistiche, trend e pattern.

        Args:
            doc: Documento Word
            results: Dizionario con i risultati delle analisi SLA
        """
        doc.add_heading('Analisi Dettagliata', level=1)

        for sla_name, data in results.items():
            doc.add_heading(f'Analisi {sla_name}', level=2)

            # Ottieni i dati giornalieri
            daily_stats = data['daily_stats']
            threshold = self.thresholds[sla_name]

            # Statistiche descrittive
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistiche Descrittive:').bold = True
            stats_para.add_run('\n')

            # Calcola statistiche di base
            total_ops = daily_stats['Time'].sum()
            over_threshold = daily_stats[sla_name].sum()
            compliance = ((total_ops - over_threshold) / total_ops * 100) if total_ops > 0 else 0

            # Aggiungi statistiche al paragrafo
            stats_table = doc.add_table(rows=1, cols=2)
            stats_table.style = 'Table Grid'

            # Headers
            header_cells = stats_table.rows[0].cells
            header_cells[0].text = 'Metrica'
            header_cells[1].text = 'Valore'

            # Dati
            metrics = [
                ('Totale Operazioni', f"{total_ops:,}"),
                ('Operazioni Oltre Soglia', f"{over_threshold:,}"),
                ('Compliance Media', f"{compliance:.2f}%"),
                ('Giorni Analizzati', str(len(daily_stats))),
                ('Soglia', f"{threshold} ms"),
                ('Worst Day', self.get_worst_day(daily_stats, sla_name)),
                ('Best Day', self.get_best_day(daily_stats, sla_name))
            ]

            for metric, value in metrics:
                cells = stats_table.add_row().cells
                cells[0].text = metric
                cells[1].text = value

            # Analisi trend
            trend_para = doc.add_paragraph()
            doc.add_paragraph()  # Spazio
            trend_para.add_run('Analisi Trend:').bold = True
            trend_para.add_run('\n')

            # Calcola trend
            trend_analysis = self.calculate_trend(daily_stats, sla_name)

            trend_para.add_run(f"• Tendenza generale: {trend_analysis['direction']}\n")
            trend_para.add_run(f"• Variazione media giornaliera: {trend_analysis['daily_change']:.2f}%\n")
            trend_para.add_run(f"• Stabilità: {trend_analysis['stability']}\n")

            if trend_analysis['warning']:
                warning_run = trend_para.add_run(f"⚠ {trend_analysis['warning']}")
                warning_run.font.color.rgb = RGBColor(255, 0, 0)

            # Pattern ricorrenti
            patterns_para = doc.add_paragraph()
            doc.add_paragraph()  # Spazio
            patterns_para.add_run('Pattern Ricorrenti:').bold = True
            patterns_para.add_run('\n')

            # Analizza pattern
            patterns = self.analyze_patterns(daily_stats, sla_name)

            for pattern in patterns:
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.add_run(pattern)

            doc.add_paragraph()  # Spazio tra sezioni

        def get_worst_day(self, daily_stats, sla_name):
            """Trova il giorno con la performance peggiore."""
            worst_day = daily_stats.loc[daily_stats[sla_name].idxmax()]
            return f"{worst_day['Date'].strftime('%Y-%m-%d')} ({worst_day[sla_name]} operazioni oltre soglia)"

    def get_best_day(self, daily_stats, sla_name):
            """Trova il giorno con la performance migliore."""
            best_day = daily_stats.loc[daily_stats[sla_name].idxmin()]
            return f"{best_day['Date'].strftime('%Y-%m-%d')} ({best_day[sla_name]} operazioni oltre soglia)"

    def calculate_trend(self, daily_stats, sla_name):
            """
            Calcola il trend dei dati.
            """
            try:
                # Calcola variazioni giornaliere
                daily_changes = daily_stats[sla_name].pct_change()
                avg_change = daily_changes.mean()

                # Determina direzione del trend
                if abs(avg_change) < 0.01:
                    direction = "Stabile"
                else:
                    direction = "In miglioramento" if avg_change < 0 else "In peggioramento"

                # Calcola stabilità
                stability_score = daily_changes.std()
                if stability_score < 0.1:
                    stability = "Molto stabile"
                elif stability_score < 0.2:
                    stability = "Moderatamente stabile"
                else:
                    stability = "Instabile"

                # Genera warning se necessario
                warning = ""
                if direction == "In peggioramento" and stability_score > 0.2:
                    warning = "Trend negativo significativo rilevato - richiede attenzione"

                return {
                    'direction': direction,
                    'daily_change': avg_change * 100,
                    'stability': stability,
                    'warning': warning
                }

            except Exception as e:
                self.log(f"Errore nel calcolo del trend per {sla_name}: {str(e)}", "error")
                return {
                    'direction': "Non calcolabile",
                    'daily_change': 0,
                    'stability': "Non determinabile",
                    'warning': "Errore nell'analisi del trend"
                }

    def analyze_patterns(self, daily_stats, sla_name):
            """
            Analizza i pattern ricorrenti nei dati.
            """
            patterns = []
            try:
                # Analisi giorni della settimana
                daily_stats['DayOfWeek'] = pd.to_datetime(daily_stats['Date']).dt.day_name()
                weekday_stats = daily_stats.groupby('DayOfWeek')[sla_name].mean()
                worst_day = weekday_stats.idxmax()
                best_day = weekday_stats.idxmin()

                patterns.append(f"Il giorno della settimana con performance peggiore è {worst_day}")
                patterns.append(f"Il giorno della settimana con performance migliore è {best_day}")

                # Analisi inizio/fine mese
                daily_stats['DayOfMonth'] = pd.to_datetime(daily_stats['Date']).dt.day
                early_month = daily_stats[daily_stats['DayOfMonth'] <= 5][sla_name].mean()
                late_month = daily_stats[daily_stats['DayOfMonth'] >= 25][sla_name].mean()

                if abs(early_month - late_month) > daily_stats[sla_name].std():
                    if early_month > late_month:
                        patterns.append("Performance tendenzialmente peggiore a inizio mese")
                    else:
                        patterns.append("Performance tendenzialmente peggiore a fine mese")

                # Identifica sequenze di peggioramento
                deterioration_seq = (daily_stats[sla_name] > daily_stats[sla_name].shift()).rolling(window=3).sum()
                if (deterioration_seq >= 3).any():
                    patterns.append("Rilevate sequenze di peggioramento consecutivo (3+ giorni)")

                # Analisi variabilità
                daily_var = daily_stats[sla_name].std()
                if daily_var > daily_stats[sla_name].mean() * 0.5:
                    patterns.append("Alta variabilità giornaliera nelle performance")

            except Exception as e:
                self.log(f"Errore nell'analisi dei pattern per {sla_name}: {str(e)}", "error")
                patterns = ["Impossibile determinare i pattern - errore nell'analisi"]

            return patterns
    def generate_docx_report(self, all_results, output_path):
        """Genera il report Word (.docx) finale."""
        try:
            doc = Document()

            # Stile del documento
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            # Prima genera le sezioni SLA1
            self.generate_sla1_section(doc, is_incoming=True)
            self.generate_sla1_section(doc, is_incoming=False)

            # Poi genera le altre sezioni SLA come prima
            for sla_name in ['SLA2', 'SLA3', 'SLA4', 'SLA5', 'SLA6']:
                file_key = f'{sla_name}_{"OUT" if sla_name in ["SLA2", "SLA3"] else "IN"}'
                if self.files[file_key]['path'] is None:
                    self.log(f"File mancante per {sla_name}", "error")
                    continue

                try:
                    df = pd.read_csv(self.files[file_key]['path'], sep=';', index_col='Time')
                    df.index = pd.to_datetime(df.index, dayfirst=True)
                    results = self.analyze_sla_data(df, self.thresholds[sla_name], sla_name)

                    section_num = int(sla_name.replace('SLA', ''))
                    doc.add_heading(f'3.{section_num} Rilevazioni {sla_name}', level=2)
                    doc.add_paragraph()

                    # Creazione tabella
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'

                    # Intestazioni
                    headers = [
                        'Data',
                        f'Min con Operatività oltre\nsoglia {self.thresholds[sla_name]} ms',
                        'Min con\nOperatività',
                        '%'
                    ]
                    for i, text in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = text
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Dati giornalieri
                    daily_stats = results['daily_stats']
                    sla_col = sla_name
                    for _, row in daily_stats.iterrows():
                        row_cells = table.add_row().cells
                        row_cells[0].text = row['Date'].strftime('%Y-%m-%d')
                        row_cells[1].text = str(int(row[sla_col]))
                        row_cells[2].text = str(int(row['Time']))

                        total_minutes = int(row['Time'])
                        over_threshold = int(row[sla_col])
                        compliance = (
                            (total_minutes - over_threshold) / total_minutes * 100
                        ) if total_minutes > 0 else 0
                        row_cells[3].text = f"{compliance:.2f}"

                        for cell in row_cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Riga RILEVAZIONE
                    row_cells = table.add_row().cells
                    row_cells[0].text = "RILEVAZIONE"
                    row_cells[1].text = str(results['total_over_threshold'])
                    row_cells[2].text = str(results['total_records'])

                    total_compliance = (
                        (results['total_records'] - results['total_over_threshold']) / results['total_records'] * 100
                    ) if results['total_records'] > 0 else 0
                    row_cells[3].text = f"{total_compliance:.2f}"

                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Conclusioni
                    doc.add_paragraph()
                    threshold_val = 95
                    is_compliant = total_compliance >= threshold_val

                    p = doc.add_paragraph()
                    p.add_run("Il ").italic = True
                    p.add_run("Service Level Agreement").italic = True
                    p.add_run(" si ritiene ")
                    if not is_compliant:
                        p.add_run("non ")
                    p.add_run("rispettato nel periodo di riferimento (")
                    bold_perc = p.add_run(f"{total_compliance:.2f}%")
                    bold_perc.bold = True
                    p.add_run(").")

                    doc.add_paragraph(f"Ciò considerato l'obiettivo previsto del {threshold_val}%.")
                    doc.add_paragraph()

                except Exception as e:
                    self.log(f"Errore nell'elaborazione di {sla_name}: {str(e)}", "error")
                    raise

            doc.save(output_path)
            return True

        except Exception as e:
            self.log(f"Errore nella generazione del report Word: {str(e)}", "error")
            raise

    def generate_sla1_section(self, doc, is_incoming=True):
        """
        Genera la sezione del report per SLA1 (Incoming o Outgoing).
        """
        try:
            sla_type = "Incoming" if is_incoming else "Outgoing"
            file_key = f'SLA1_{"IN" if is_incoming else "OUT"}'

            if self.files[file_key]['path'] is None:
                self.log(f"File mancante per SLA1 {sla_type}", "error")
                return

            # Carica e analizza i dati
            df = pd.read_csv(self.files[file_key]['path'], sep=';')
            results = self.analyze_sla1_data(df, is_incoming)

            # Aggiungi intestazione
            doc.add_heading(f'3.1 Rilevazioni SLA 1 {sla_type}', level=2)
            doc.add_paragraph("Throughput Garantito Operazioni Gestite.")
            doc.add_paragraph("Gestione di un picco di 50 transazioni al minuto.")

            # Crea tabella
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Intestazioni
            headers = [
                'Data',
                'Min in cui non sono state\ngestite tutte le operazioni\nricevute (se <50)',
                'Min Totali',
                '%'
            ]

            # Formatta intestazioni
            for i, text in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = text
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Aggiungi dati
            daily_stats = results['daily_stats']
            for _, row in daily_stats.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = row['Date'].strftime('%d/%m/%Y')
                row_cells[1].text = str(int(row['Failed_Minutes']))
                row_cells[2].text = str(int(row['Total_Minutes']))
                row_cells[3].text = f"{row['Compliance']:.2f}"

                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Aggiungi riga RILEVAZIONE
            row_cells = table.add_row().cells
            row_cells[0].text = "RILEVAZIONE"
            row_cells[1].text = str(results['total_over_threshold'])
            row_cells[2].text = str(results['total_records'])

            total_compliance = ((results['total_records'] - results['total_over_threshold']) /
                                results['total_records'] * 100)
            row_cells[3].text = f"{total_compliance:.2f}"

            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Aggiungi conclusioni
            doc.add_paragraph()
            threshold_val = 95
            is_compliant = total_compliance >= threshold_val

            p = doc.add_paragraph()
            p.add_run("Il ").italic = True
            p.add_run("Service Level Agreement").italic = True
            p.add_run(" si ritiene pertanto ")
            if not is_compliant:
                p.add_run("non ")
            p.add_run(f"rispettato nel periodo di riferimento (")
            bold_perc = p.add_run(f"{total_compliance:.2f}%")
            bold_perc.bold = True
            p.add_run(").")

            doc.add_paragraph(f"Ciò considerato l'obiettivo previsto del {threshold_val}%.")
            doc.add_paragraph()

        except Exception as e:
            self.log(f"Errore nella generazione sezione SLA1 {sla_type}: {str(e)}", "error")
            raise

    def process_sla1_file(self, file_path, is_incoming=True):
        """
        Elabora il file SLA1 e calcola le statistiche necessarie.
        """
        try:
            df = pd.read_csv(file_path, sep=';')
            daily_stats = pd.DataFrame()

            # Raggruppamento per data
            df['Date'] = pd.to_datetime(df['Time']).dt.date
            daily_stats = df.groupby('Date').agg({
                'Failed_Minutes': 'sum',
                'Total_Minutes': 'first'
            }).reset_index()

            return daily_stats

        except Exception as e:
            self.log(f"Errore nell'elaborazione del file SLA1: {str(e)}", "error")
            raise

    def calculate_sla1_compliance(self, stats_df):
        """
        Calcola la conformità per SLA1.
        """
        try:
            total_minutes = stats_df['Total_Minutes'].sum()
            failed_minutes = stats_df['Failed_Minutes'].sum()
            compliance = ((total_minutes - failed_minutes) / total_minutes * 100) if total_minutes > 0 else 0
            return compliance, total_minutes, failed_minutes
        except Exception as e:
            self.log(f"Errore nel calcolo della conformità SLA1: {str(e)}", "error")
            raise

    def analyze_all_sla1(self):
        """
        Analizza entrambi i file SLA1 (Incoming e Outgoing).
        """
        results = {'IN': None, 'OUT': None}

        for direction in ['IN', 'OUT']:
            key = f'SLA1_{direction}'
            if self.files[key]['path']:
                try:
                    stats = self.process_sla1_file(self.files[key]['path'], direction == 'IN')
                    compliance, total, failed = self.calculate_sla1_compliance(stats)
                    results[direction] = {
                        'stats': stats,
                        'compliance': compliance,
                        'total_minutes': total,
                        'failed_minutes': failed
                    }
                except Exception as e:
                    self.log(f"Errore nell'analisi di {key}: {str(e)}", "error")

        return results

    def export_sla1_results(self, results, output_path):
        """
        Esporta i risultati dell'analisi SLA1 in formato CSV.
        """
        try:
            with pd.ExcelWriter(output_path) as writer:
                for direction, data in results.items():
                    if data:
                        sheet_name = f'SLA1_{direction}'
                        df = data['stats'].copy()
                        df['Compliance'] = ((df['Total_Minutes'] - df['Failed_Minutes']) /
                                            df['Total_Minutes'] * 100)
                        df.to_excel(writer, sheet_name=sheet_name, index=False)

            self.log(f"Risultati SLA1 esportati in {output_path}", "success")
        except Exception as e:
            self.log(f"Errore nell'esportazione dei risultati: {str(e)}", "error")
            raise

    def validate_sla1_data(self, df):
        """
        Valida i dati di input per SLA1.
        """
        required_columns = ['Time', 'Failed_Minutes', 'Total_Minutes']
        missing_columns = [col for col in required_columns if col not in df.columns]

        if missing_columns:
            raise ValueError(f"Colonne mancanti: {', '.join(missing_columns)}")

        # Verifica valori numerici
        for col in ['Failed_Minutes', 'Total_Minutes']:
            if not pd.to_numeric(df[col], errors='coerce').notnull().all():
                raise ValueError(f"La colonna {col} contiene valori non numerici")

        # Verifica valori logici
        if (df['Failed_Minutes'] > df['Total_Minutes']).any():
            raise ValueError("I minuti falliti non possono essere maggiori del totale")

        return True

    def add_summary_to_report(self, doc, sla1_results):
        """
        Aggiunge una sezione di riepilogo al report.
        """
        doc.add_heading('1. Riepilogo Generale', level=1)
        doc.add_paragraph()

        # Crea tabella riepilogativa
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Intestazioni
        headers = ['SLA', 'Obiettivo', 'Risultato', 'Status']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Aggiunge dati SLA1
        for direction, data in sla1_results.items():
            if data:
                row = table.add_row().cells
                row[0].text = f'SLA1 {direction}'
                row[1].text = '95%'
                row[2].text = f"{data['compliance']:.2f}%"
                row[3].text = 'OK' if data['compliance'] >= 95 else 'KO'

        doc.add_paragraph()

    def create_detailed_sla1_section(self, doc, direction, data):
        """
        Crea una sezione dettagliata per SLA1 nel report.
        """
        doc.add_heading(f'3.1 Rilevazioni SLA1 {direction}', level=2)
        doc.add_paragraph("Throughput Garantito Operazioni Gestite.")
        doc.add_paragraph("Gestione di un picco di 50 transazioni al minuto.")
        doc.add_paragraph()

        # Crea tabella dettagli
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Intestazioni
        headers = ['Data', 'Min Falliti', 'Min Totali', 'Compliance %']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Aggiunge dati giornalieri
        for _, row in data['stats'].iterrows():
            cells = table.add_row().cells
            cells[0].text = row['Date'].strftime('%d/%m/%Y')
            cells[1].text = str(int(row['Failed_Minutes']))
            cells[2].text = str(int(row['Total_Minutes']))
            compliance = ((row['Total_Minutes'] - row['Failed_Minutes']) /
                          row['Total_Minutes'] * 100)
            cells[3].text = f"{compliance:.2f}%"

        # Aggiunge riga totali
        cells = table.add_row().cells
        cells[0].text = 'TOTALE'
        cells[1].text = str(int(data['failed_minutes']))
        cells[2].text = str(int(data['total_minutes']))
        cells[3].text = f"{data['compliance']:.2f}%"
        for cell in cells:
            cell.paragraphs[0].runs[0].bold = True

    def create_error_report(self, sla1_results):
        """
        Crea un report degli errori per SLA1.
        """
        error_report = []

        for direction, data in sla1_results.items():
            if data:
                stats = data['stats']
                # Trova giorni con alta percentuale di fallimenti
                problem_days = stats[stats['Failed_Minutes'] / stats['Total_Minutes'] > 0.1]

                if not problem_days.empty:
                    error_report.append(f"\nSLA1 {direction} - Giorni critici:")
                    for _, row in problem_days.iterrows():
                        error_pct = (row['Failed_Minutes'] / row['Total_Minutes'] * 100)
                        error_report.append(
                            f"- {row['Date']}: {error_pct:.1f}% di fallimenti "
                            f"({row['Failed_Minutes']} minuti su {row['Total_Minutes']})"
                        )

        return '\n'.join(error_report)

    def configure_logging(self):
        """
        Configura il sistema di logging avanzato.
        """
        import logging

        # Configura il logger di base
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('sla_analysis.log'),
                logging.StreamHandler()
            ]
        )

        # Crea logger specifico per SLA
        self.logger = logging.getLogger('SLA_Analysis')
        self.logger.setLevel(logging.INFO)

        # Aggiungi handler per la GUI
        class GUIHandler(logging.Handler):
            def __init__(self, text_widget):
                super().__init__()
                self.text_widget = text_widget

            def emit(self, record):
                msg = self.format(record)
                self.text_widget.insert('end', msg + '\n')
                self.text_widget.see('end')

        gui_handler = GUIHandler(self.log_text)
        gui_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s: %(message)s'))
        self.logger.addHandler(gui_handler)


    def create_backup(self):
        """
        Crea un backup dei file di input.
        """
        backup_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "backup")
        os.makedirs(backup_dir, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        for key, file_info in self.files.items():
            if file_info['path']:
                try:
                    src = file_info['path']
                    dst = os.path.join(backup_dir, f"{key}_{timestamp}.csv")
                    shutil.copy2(src, dst)
                    self.log(f"Backup creato per {key}: {dst}", "info")
                except Exception as e:
                    self.log(f"Errore nel backup di {key}: {str(e)}", "error")

    def cleanup_old_backups(self, days=30):
        """
        Pulisce i backup più vecchi di X giorni.
        """
        backup_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "backup")
        if not os.path.exists(backup_dir):
            return

        cutoff = datetime.now() - timedelta(days=days)

        for file in os.listdir(backup_dir):
            file_path = os.path.join(backup_dir, file)
            if os.path.isfile(file_path):
                file_time = datetime.fromtimestamp(os.path.getctime(file_path))
                if file_time < cutoff:
                    try:
                        os.remove(file_path)
                        self.log(f"Rimosso backup obsoleto: {file}", "info")
                    except Exception as e:
                        self.log(f"Errore nella rimozione del backup {file}: {str(e)}", "error")

    def manage_backups(self):
        """
        Gestisce manualmente i backup.
        """
        backup_window = tk.Toplevel(self)
        backup_window.title("Gestione Backup")
        backup_window.geometry("600x400")

        # Lista dei backup
        backup_frame = ttk.Frame(backup_window, padding=10)
        backup_frame.pack(fill='both', expand=True)

        # Treeview per i backup
        columns = ('File', 'Data', 'Dimensione')
        tree = ttk.Treeview(backup_frame, columns=columns, show='headings')

        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=100)

        # Scrollbar
        scrollbar = ttk.Scrollbar(backup_frame, orient='vertical', command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)

        # Bottoni
        btn_frame = ttk.Frame(backup_window, padding=10)
        ttk.Button(btn_frame, text="Elimina Selezionati",
                   command=lambda: self.delete_selected_backups(tree)).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="Backup Manuale",
                   command=self.create_backup).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="Aggiorna Lista",
                   command=lambda: self.update_backup_list(tree)).pack(side='left', padx=5)

        # Layout
        tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        btn_frame.pack(fill='x')

        # Popola la lista
        self.update_backup_list(tree)

    def update_backup_list(self, tree):
        """
        Aggiorna la lista dei backup nel treeview.
        """
        # Pulisci la lista esistente
        for item in tree.get_children():
            tree.delete(item)

        # Recupera la lista dei backup
        backup_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "backup")
        if not os.path.exists(backup_dir):
            return

        for file in os.listdir(backup_dir):
            file_path = os.path.join(backup_dir, file)
            if os.path.isfile(file_path):
                # Ottieni informazioni sul file
                stats = os.stat(file_path)
                size = stats.st_size / 1024  # KB
                date = datetime.fromtimestamp(stats.st_mtime)

                tree.insert('', 'end', values=(
                    file,
                    date.strftime('%Y-%m-%d %H:%M:%S'),
                    f"{size:.1f} KB"
                ))

    def delete_selected_backups(self, tree):
        """
        Elimina i backup selezionati con gestione degli errori avanzata e feedback all'utente.

        Args:
            tree: Il Treeview contenente la lista dei backup
        """
        # Verifica selezione
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Attenzione", "Nessun backup selezionato")
            return

        # Conteggio file selezionati
        num_selected = len(selected)
        if not messagebox.askyesno("Conferma",
                                   f"Eliminare {num_selected} backup selezionati?"):
            return

        backup_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "backup")

        # Contatori per il report finale
        success_count = 0
        error_count = 0
        error_files = []

        try:
            # Verifica che la directory backup esista
            if not os.path.exists(backup_dir):
                raise FileNotFoundError("Directory backup non trovata")

            # Processa ogni file selezionato
            for item in selected:
                values = tree.item(item)['values']
                if not values:
                    continue

                file_name = values[0]
                file_path = os.path.join(backup_dir, file_name)

                try:
                    # Verifica che il file esista ancora
                    if not os.path.exists(file_path):
                        raise FileNotFoundError(f"File {file_name} non trovato")

                    # Verifica permessi prima di eliminare
                    if not os.access(file_path, os.W_OK):
                        raise PermissionError(f"Permessi insufficienti per {file_name}")

                    # Elimina il file
                    os.remove(file_path)
                    success_count += 1
                    self.log(f"Backup eliminato manualmente: {file_name}", "success")

                except (FileNotFoundError, PermissionError) as e:
                    error_count += 1
                    error_files.append(file_name)
                    self.log(f"Errore nell'eliminazione del backup {file_name}: {str(e)}", "error")
                    continue

            # Aggiorna la lista dei backup
            self.update_backup_list(tree)

            # Prepara il messaggio di report finale
            report_msg = f"Eliminati con successo: {success_count}/{num_selected} backup"
            if error_count > 0:
                report_msg += f"\nErrori: {error_count}"
                if error_files:
                    report_msg += f"\nFile con errori:\n" + "\n".join(f"- {f}" for f in error_files)

            # Mostra il report appropriato
            if error_count == 0:
                messagebox.showinfo("Completato", report_msg)
            else:
                messagebox.showwarning("Completato con errori", report_msg)

        except Exception as e:
            # Gestione errori generali
            error_msg = f"Errore durante l'eliminazione dei backup:\n{str(e)}"
            self.log(error_msg, "error")
            messagebox.showerror("Errore", error_msg)

        finally:
            # Aggiorna sempre la lista alla fine
            try:
                self.update_backup_list(tree)
            except Exception as e:
                self.log(f"Errore nell'aggiornamento della lista: {str(e)}", "error")

    def set_cell_border(self, tc, top=None, bottom=None):
        """Imposta i bordi di una cella della tabella."""
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

        if top:
            top_element = OxmlElement('w:top')
            top_element.set(qn('w:sz'), str(top['sz']))
            top_element.set(qn('w:val'), top['val'])
            top_element.set(qn('w:color'), top['color'].replace('#', ''))
            tcBorders.append(top_element)

        if bottom:
            bottom_element = OxmlElement('w:bottom')
            bottom_element.set(qn('w:sz'), str(bottom['sz']))
            bottom_element.set(qn('w:val'), bottom['val'])
            bottom_element.set(qn('w:color'), bottom['color'].replace('#', ''))
            tcBorders.append(bottom_element)

    def fast_normalize_time_value(self, value):
        """
        Converte una stringa di tempo in millisecondi (float).
        Versione senza log a ogni riga, per maggiore velocità.
        """
        import re
        import numpy as np

        if pd.isna(value):
            return 0.0  # O np.nan, a scelta

        v = str(value).strip().lower()

        # 1) Se è già un numero (es. "1200" o "1200.5")
        try:
            return float(v)
        except ValueError:
            pass

        # 2) millisecondi con "ms" (es. "647 ms")
        if v.endswith('ms'):
            numeric_str = v.replace('ms', '').strip()
            return float(numeric_str)

        # 3) secondi con "s" (es. "10s")
        if re.match(r'^\d+(\.\d+)?\s*s$', v):
            numeric_str = re.sub(r's$', '', v).strip()
            return float(numeric_str) * 1000.0

        # 4) minuti (es. "2 min", "2mins", "2minutes")
        if re.search(r'(?:mins?|minutes?)$', v):
            numeric_str = re.sub(r'(?:mins?|minutes?)$', '', v).strip()
            return float(numeric_str) * 60000.0

        # Se non riconosciuto, restituiamo 0.0 (o np.nan)
        return 0.0

    def analyze_sla_data(self, df, threshold, sla_name):
        """
        Analizza i dati SLA per produrre statistiche giornaliere in modo più veloce.
        """
        try:
            # Copia DF
            df = df.copy()

            # Se Time è nell'indice, lo resettiamo come colonna
            if df.index.name == 'Time':
                df = df.reset_index()

            # Verifica colonne
            required_columns = ['Time', sla_name]
            for col in required_columns:
                if col not in df.columns:
                    raise KeyError(f"Colonna {col} non trovata nel DataFrame per {sla_name}")

            # Conversione colonna Time in datetime
            df['Time'] = pd.to_datetime(df['Time'], dayfirst=True, errors='coerce')
            if df['Time'].isnull().any():
                self.log(f"Attenzione: ci sono date non parse correttamente per {sla_name}.", "warning")

            # Conversione vettoriale della colonna SLA
            df[sla_name] = df[sla_name].apply(self.fast_normalize_time_value)

            # Aggiunge colonna Date
            df['Date'] = df['Time'].dt.date

            # Calcolo statistiche giornaliere
            daily_stats = df.groupby('Date').agg({
                sla_name: lambda x: (x > threshold).sum(),
                'Time': 'count'
            }).reset_index()

            # Totali
            total_records = int(daily_stats['Time'].sum())
            total_over_threshold = int(daily_stats[sla_name].sum())

            # Dettagli di non conformità (vettoriale)
            non_compliant_df = df[df[sla_name] > threshold].copy()
            non_compliant_df['deviation'] = non_compliant_df[sla_name] - threshold
            non_compliant_details = non_compliant_df[['Time', sla_name, 'deviation']] \
                .rename(columns={sla_name: 'value'}) \
                .to_dict(orient='records')

            results = {
                'daily_stats': daily_stats,
                'total_records': total_records,
                'total_over_threshold': total_over_threshold,
                'non_compliant_details': non_compliant_details,
                'threshold': threshold
            }

            # Log finale di riepilogo
            self.log(f"[{sla_name}] Totale record: {total_records}, Oltre soglia: {total_over_threshold}", "info")
            return results

        except Exception as e:
            self.log(f"Errore nell'analisi di {sla_name}: {str(e)}", "error")
            raise

    def format_table_row(self, row_cells, bold=False, center=True):
        """Formatta una riga della tabella con le impostazioni specificate."""
        for cell in row_cells:
            if center:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if bold:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    def add_table_borders(self, table):
        """Aggiunge i bordi alla tabella."""
        tbl = table._tbl
        borders = tbl.xpath(".//w:tcBorders")
        for border in borders:
            for side in ['top', 'bottom', 'left', 'right']:
                side_element = OxmlElement(f'w:{side}')
                side_element.set(qn('w:sz'), '4')
                side_element.set(qn('w:val'), 'single')
                side_element.set(qn('w:color'), '000000')
                border.append(side_element)

    def process_sla_file(self, sla_name, file_path):
        """Elabora un file SLA e restituisce i dati analizzati."""
        try:
            df = pd.read_csv(file_path, sep=';')
            threshold = self.thresholds[sla_name]
            results = self.analyze_sla_data(df, threshold, sla_name)
            return results
        except Exception as e:
            self.log(f"Errore nell'elaborazione del file {sla_name}: {str(e)}", "error")
            raise

    def add_summary_section(self, doc, total_compliance):
        """Aggiunge la sezione di riepilogo al documento."""
        doc.add_paragraph()
        summary = doc.add_heading('Riepilogo Generale', level=1)
        summary.style = doc.styles['Heading 1']

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Indicatore'
        header_cells[1].text = 'Valore'

        row_cells = table.add_row().cells
        row_cells[0].text = 'Compliance Media'
        row_cells[1].text = f"{total_compliance:.2f}%"

        self.format_table_row(header_cells, bold=True)
        self.format_table_row(row_cells)

    def check_required_files(self):
        """Verifica che tutti i file necessari siano stati selezionati."""
        missing = []
        for file_key, file_data in self.files.items():
            if file_data['path'] is None:
                missing.append(file_key)
        return missing

    def validate_month_year(self):
        """Valida il mese e l'anno selezionati."""
        try:
            month = self.month_var.get()
            year = int(self.year_var.get())
            if year < 2000 or year > 2100:
                raise ValueError("Anno non valido")
            return True
        except:
            return False

    def search(self, text):
        """Ricerca testuale nei file SLA caricati e nel log."""
        results = []
        search_text = text.lower()

        # Ricerca nel log
        log_content = self.log_text.get(1.0, tk.END).lower()
        if search_text in log_content:
            relevant_lines = [line for line in log_content.split('\n') if search_text in line.lower()]
            if relevant_lines:
                results.append(("Log SLA", "\n".join(relevant_lines)))

        # Ricerca nei file caricati
        for file_key, file_data in self.files.items():
            if file_data['path']:
                try:
                    df = pd.read_csv(file_data['path'], sep=';')
                    for column in df.columns:
                        matches = df[df[column].astype(str).str.contains(search_text, case=False, na=False)]
                        if not matches.empty:
                            results.append((
                                f"File {file_key} - Colonna {column}",
                                f"Trovate {len(matches)} corrispondenze\n" + matches.head().to_string()
                            ))
                except Exception as e:
                    self.log(f"Errore nella ricerca nel file {file_key}: {str(e)}", "error")

        # Ricerca nelle configurazioni SLA
        try:
            for sla_name, threshold in self.thresholds.items():
                if search_text in sla_name.lower():
                    results.append((
                        "Configurazione SLA",
                        f"SLA: {sla_name}\nSoglia: {threshold}"
                    ))
        except Exception as e:
            self.log(f"Errore nella ricerca configurazioni SLA: {str(e)}", "error")

        if not results:
            results.append(("Info", f"Nessun risultato trovato per: {text}"))

        self.log(f"Ricerca eseguita per: {text} - Trovati {len(results)} risultati", "info")

        return results


def main():
    root = tk.Tk()
    root.title("SLA Analysis")
    app = SLATab(root, current_user="Utente", user_role="admin")
    app.pack(fill='both', expand=True)
    root.mainloop()


if __name__ == "__main__":
    main()
